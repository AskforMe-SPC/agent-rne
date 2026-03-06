import io
import html
import json
import logging
import os
import re
import sys
import threading
import time
import uuid
import zipfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

import requests
from flask import Flask, jsonify, render_template, request, send_file
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import mm, cm
from reportlab.lib.colors import HexColor, white
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    BaseDocTemplate, Frame, PageTemplate, NextPageTemplate,
    Paragraph as RLPara, Spacer, Table, TableStyle, PageBreak,
    Image as RLImage, HRFlowable, Flowable,
)
from docx import Document as DocxDocument
from docx.shared import Pt, Cm as DocxCm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

try:
    from dotenv import load_dotenv
except Exception:
    def load_dotenv(*args: Any, **kwargs: Any) -> bool:
        return False

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from openpyxl import load_workbook
except Exception:
    load_workbook = None

BASE_DIR = Path(__file__).resolve().parent
load_dotenv(dotenv_path=BASE_DIR / ".env", override=True)
ASSETS_DIR = BASE_DIR / "assets"
LOGO_PATH = ASSETS_DIR / "logo.png"

INPI_BASE = os.getenv("INPI_BASE_URL", "https://registre-national-entreprises.inpi.fr/api")
INPI_USERNAME = os.getenv("INPI_USERNAME", "ask.me.spc@gmail.com")
INPI_PASSWORD = os.getenv("INPI_PASSWORD", "5Z-36Cg+g$*LJw.")
PORT = int(os.getenv("PORT", "5000"))

TYPE_DOC_XLSX = os.getenv(
    "INPI_NOMENCLATURE_XLSX",
    str(Path(__file__).parent / "- ressources -" / "Dictionnaire_de_donnees_INPI_2025_05_09.xlsx"),
)
LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()
logging.basicConfig(level=LOG_LEVEL, format="%(asctime)s | %(levelname)s | %(message)s")
logger = logging.getLogger("agent-rne-vscode")

app = Flask(__name__, template_folder=str(BASE_DIR / "templates"))


@dataclass
class Job:
    siren: str
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:8])
    status: str = "pending"
    progress: int = 0
    steps: List[Dict[str, Any]] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    denomination: str = ""
    bilans_count: int = 0
    actes_count: int = 0
    total_docs: int = 0
    estimated_minutes: int = 0
    current_doc_name: str = ""
    current_doc_desc: str = ""
    zip_data: Optional[bytes] = None

    def log(self, message: str, progress: Optional[int] = None) -> None:
        if progress is not None:
            self.progress = progress
        self.steps.append(
            {
                "time": datetime.now().isoformat(timespec="seconds"),
                "message": message,
                "progress": self.progress,
            }
        )
        logger.info("[%s] %s", self.id, message)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "siren": self.siren,
            "status": self.status,
            "progress": self.progress,
            "steps": self.steps,
            "errors": self.errors,
            "denomination": self.denomination,
            "bilans_count": self.bilans_count,
            "actes_count": self.actes_count,
            "total_docs": self.total_docs,
            "estimated_minutes": self.estimated_minutes,
            "current_doc_name": self.current_doc_name,
            "current_doc_desc": self.current_doc_desc,
        }


jobs: Dict[str, Job] = {}
jobs_lock = threading.Lock()


def make_http_session() -> requests.Session:
    retry = Retry(
        total=4,
        backoff_factor=0.8,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["GET", "POST"],
        raise_on_status=False,
    )
    adapter = HTTPAdapter(max_retries=retry)
    s = requests.Session()
    s.mount("https://", adapter)
    s.mount("http://", adapter)
    return s


HTTP = make_http_session()


def validate_config() -> tuple[bool, List[str]]:
    missing: List[str] = []
    if not INPI_USERNAME.strip():
        missing.append("INPI_USERNAME")
    if not INPI_PASSWORD.strip():
        missing.append("INPI_PASSWORD")
    return (len(missing) == 0, missing)


def sanitize_filename(value: str, max_len: int = 180) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|;]", "", value)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    cleaned = cleaned.rstrip(".")
    if not cleaned:
        cleaned = "document"
    return cleaned[:max_len]


def valid_date(*candidates: Any) -> str:
    for c in candidates:
        if not c:
            continue
        s = str(c)
        m = re.search(r"(\d{4})-(\d{2})-(\d{2})", s)
        if m:
            return f"{m.group(1)}-{m.group(2)}-{m.group(3)}"
    return ""


def format_date_fr(value: Any) -> str:
    s = str(value or "").strip()
    if not s:
        return "-"
    m = re.search(r"(\d{4})-(\d{2})-(\d{2})", s)
    if m:
        return f"{m.group(3)}/{m.group(2)}/{m.group(1)}"
    return s


def format_date_fr_optional(value: Any) -> str:
    s = str(value or "").strip()
    if not s:
        return ""
    return format_date_fr(s)


def as_dash(value: Any) -> str:
    s = str(value or "").strip()
    return s if s else "-"


# ══════════════════════════════════════════════════════════════════════════════
# NOMENCLATURE INPI — chargement unique au démarrage
# ══════════════════════════════════════════════════════════════════════════════

def _load_sheet_map(wb: Any, sheet: str) -> Dict[str, str]:
    """Charge un onglet code→libellé depuis un workbook openpyxl ouvert."""
    if sheet not in wb.sheetnames:
        return {}
    ws = wb[sheet]
    out: Dict[str, str] = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row or row[0] is None:
            continue
        code = str(row[0]).strip()
        label = str(row[1]).strip() if len(row) > 1 and row[1] is not None else ""
        if code and label:
            out[code] = label
    return out


def load_inpi_nomenclature(xlsx_path: str) -> Dict[str, Dict[str, str]]:
    """Charge les onglets utiles du dictionnaire INPI en une seule passe."""
    if not load_workbook:
        logger.warning("openpyxl non disponible → nomenclature vide")
        return {}
    p = Path(xlsx_path)
    if not p.exists():
        logger.warning("Fichier nomenclature INPI introuvable: %s", p)
        return {}
    try:
        wb = load_workbook(filename=str(p), read_only=True, data_only=True)
        target_sheets = [
            "typeDocument", "formeJuridique", "roleEntreprise", "typeVoie",
            "typeDePersonne", "deviseCapital", "exerciceActivite", "formeExercice",
            "typeOrigine", "pays", "diffusionINSEE", "typeActe", "natureActe",
        ]
        maps: Dict[str, Dict[str, str]] = {}
        for s in target_sheets:
            m = _load_sheet_map(wb, s)
            if m:
                maps[s] = m
                logger.info("Nomenclature '%s' chargée: %s entrées", s, len(m))
        wb.close()
        return maps
    except Exception as e:
        logger.error("Erreur chargement nomenclature INPI: %s", e)
        return {}


INPI_MAPS: Dict[str, Dict[str, str]] = load_inpi_nomenclature(TYPE_DOC_XLSX)
TYPE_DOC_MAP: Dict[str, str] = INPI_MAPS.get("typeDocument", {})


def resolve_label(category: str, code: Any, fallback: str = "") -> str:
    """Résout un code INPI en libellé lisible. Ne plante jamais."""
    if code is None:
        return fallback
    code_str = str(code).strip()
    if not code_str:
        return fallback
    m = INPI_MAPS.get(category, {})
    label = m.get(code_str) or m.get(code_str.upper()) or m.get(code_str.lower())
    if not label and code_str.isdigit():
        label = m.get(str(int(code_str)))
    return label or fallback or code_str


def resolve_label_multi(categories: List[str], code: Any, fallback: str = "") -> str:
    code_str = str(code or "").strip()
    if not code_str:
        return fallback
    for category in categories:
        lbl = resolve_label(category, code_str, fallback="")
        if lbl and lbl != code_str:
            return lbl
    return fallback or code_str


def resolve_person_type_label(code: Any) -> str:
    """Convertit les codes courts INPI de type personne en libellés lisibles."""
    c = str(code or "").strip().upper()
    if not c:
        return ""
    # 1) tentative nomenclature officielle
    lbl = resolve_label_multi(["typeDePersonne"], c, fallback="")
    if lbl and lbl.upper() != c:
        return lbl
    # 2) fallback robuste sur codes courts fréquemment rencontrés
    short_map = {
        "M": "Personne morale",
        "P": "Personne physique",
        "PM": "Personne morale",
        "PP": "Personne physique",
    }
    return short_map.get(c, c)


def format_birth_month_year(value: Any) -> str:
    s = str(value or "").strip()
    if not s:
        return "-"
    m = re.search(r"^(\d{4})-(\d{2})", s)
    if m:
        return f"{m.group(2)}/{m.group(1)}"
    return s


def normalize_text(value: Any) -> str:
    s = str(value or "").strip()
    if not s:
        return ""
    # Repairs common mojibake when UTF-8 text was decoded as latin-1.
    if "Ã" in s or "Â" in s:
        try:
            return s.encode("latin-1").decode("utf-8")
        except Exception:
            return s
    return s


def type_doc_label(code: str) -> str:
    return resolve_label("typeDocument", code, fallback=code or "N/A")


def resolve_doc_type(meta: Dict[str, Any], famille: str) -> Dict[str, str]:
    code = str(meta.get("typeDocument") or "").strip().upper()
    if code:
        mapped = type_doc_label(code)
        if mapped and mapped != code:
            return {"type_code": code, "type_label": mapped}

    # Many INPI actes do not expose typeDocument and instead provide typeRdd/typeActe.
    rdd = meta.get("typeRdd")
    if isinstance(rdd, list):
        labels: List[str] = []
        seen = set()
        for row in rdd:
            if not isinstance(row, dict):
                continue
            t_code = str(row.get("typeActe") or "").strip()
            t = resolve_label_multi(["typeActe", "natureActe", "typeDocument"], t_code, fallback=t_code)
            if not t:
                continue
            k = t.casefold()
            if k not in seen:
                seen.add(k)
                labels.append(t)
        if labels:
            return {"type_code": code, "type_label": " + ".join(labels[:2])}

    if famille.upper() == "ACTE":
        for k in ["typeActe", "natureActe", "libelleTypeDocument"]:
            v = str(first_value(meta, [k]) or "").strip()
            if v:
                return {
                    "type_code": code,
                    "type_label": resolve_label_multi(
                        ["typeActe", "natureActe", "typeDocument"],
                        v,
                        fallback=v,
                    ),
                }

    return {"type_code": code, "type_label": type_doc_label(code)}


def inpi_login() -> str:
    ok, missing = validate_config()
    if not ok:
        raise ValueError(f"Configuration INPI manquante: {', '.join(missing)}")

    r = HTTP.post(
        f"{INPI_BASE}/sso/login",
        json={"username": INPI_USERNAME, "password": INPI_PASSWORD},
        headers={"Content-Type": "application/json"},
        timeout=30,
    )
    r.raise_for_status()
    token = r.json().get("token")
    if not token:
        raise ValueError("Token INPI non reçu")
    return token


def inpi_get_attachments(siren: str, token: str) -> Dict[str, Any]:
    r = HTTP.get(
        f"{INPI_BASE}/companies/{siren}/attachments",
        headers={"Authorization": f"Bearer {token}"},
        timeout=60,
    )
    r.raise_for_status()
    return r.json()


def inpi_get_company(siren: str, token: str) -> Dict[str, Any]:
    r = HTTP.get(
        f"{INPI_BASE}/companies/{siren}",
        headers={"Authorization": f"Bearer {token}"},
        timeout=60,
    )
    r.raise_for_status()
    return r.json()


def inpi_get_acte_info(acte_id: str, token: str) -> Dict[str, Any]:
    try:
        r = HTTP.get(
            f"{INPI_BASE}/actes/{acte_id}",
            headers={"Authorization": f"Bearer {token}"},
            timeout=30,
        )
        r.raise_for_status()
        return r.json().get("data", {})
    except Exception:
        return {}


def inpi_download(endpoint: str, token: str) -> bytes:
    r = HTTP.get(
        f"{INPI_BASE}/{endpoint}",
        headers={"Authorization": f"Bearer {token}"},
        timeout=90,
    )
    r.raise_for_status()
    return r.content


def inpi_download_company_export_pdf(siren: str) -> bytes:
    headers = {
        "Accept": "application/pdf,application/octet-stream;q=0.9,*/*;q=0.8",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) Agent-RNE/1.0",
        "Referer": "https://data.inpi.fr/",
    }
    attempts = [
        (f"https://data.inpi.fr/export/companies?format=pdf&ids=[%22{siren}%22]", None),
        ("https://data.inpi.fr/export/companies", {"format": "pdf", "ids": f"[\"{siren}\"]"}),
        ("https://data.inpi.fr/export/companies", {"format": "pdf", "ids": f"[{siren}]"}),
        ("https://data.inpi.fr/export/companies", {"format": "pdf", "ids": siren}),
        (f"https://data.inpi.fr/export/companies?format=pdf&ids=[{siren}]", None),
    ]
    last_err = ""
    for url, params in attempts:
        try:
            r = HTTP.get(url, params=params, headers=headers, timeout=90, allow_redirects=True)
            ctype = (r.headers.get("content-type") or "").lower()
            if r.status_code >= 400:
                last_err = f"HTTP {r.status_code}"
                continue
            if r.content and (r.content[:4] == b"%PDF" or "application/pdf" in ctype):
                return r.content
            last_err = f"Réponse non-PDF ({ctype or 'content-type inconnu'})"
        except Exception as e:
            last_err = str(e)
    raise ValueError(f"Export INPI PDF indisponible pour {siren}: {last_err}")


def extract_pdf_text(pdf_bytes: bytes) -> str:
    if not pdf_bytes:
        return ""

    if fitz is not None:
        try:
            with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
                text = "\n\n".join((page.get_text("text") or "").strip() for page in doc)
                text = text.strip()
                if text:
                    return text
        except Exception:
            pass

    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        parts = [(p.extract_text() or "").strip() for p in reader.pages]
        return "\n\n".join([p for p in parts if p]).strip()
    except Exception:
        return ""


def first_lines(text: str, max_lines: int = 6) -> str:
    cleaned: List[str] = []
    for raw in text.splitlines():
        line = re.sub(r"\s+", " ", raw).strip()
        if not line:
            continue
        if re.match(r"^page\s+\d+(\s*/\s*\d+)?$", line, flags=re.IGNORECASE):
            continue
        if re.match(r"^(rcs|greffe|www\.|http)", line, flags=re.IGNORECASE):
            continue
        cleaned.append(line)

    out: List[str] = []
    seen: Dict[str, int] = {}
    for line in cleaned:
        key = line.casefold()
        seen[key] = seen.get(key, 0) + 1
        if seen[key] > 1 and len(line) < 90:
            continue
        out.append(line)
        if len(out) >= max_lines:
            break
    return " | ".join(out)[:500]


def extract_type_rdd_rows(meta: Dict[str, Any]) -> List[Dict[str, str]]:
    out: List[Dict[str, str]] = []
    rdd = meta.get("typeRdd")
    if not isinstance(rdd, list):
        return out
    for row in rdd:
        if not isinstance(row, dict):
            continue
        label_code = str(row.get("typeActe") or "").strip()
        label = resolve_label_multi(["typeActe", "natureActe", "typeDocument"], label_code, fallback=label_code)
        decision = str(row.get("decision") or "").strip()
        if label or decision:
            out.append({"type_label": label or "N/A", "decision": decision})
    return out


def extract_reference_year(meta: Dict[str, Any]) -> str:
    """Retourne une année de référence pour les bilans (YYYY) sans lever d'exception."""
    if not isinstance(meta, dict):
        return ""

    candidates: List[Any] = []
    for key in [
        "anneeReference",
        "anneeExercice",
        "annee",
        "exercice",
        "dateClotureExercice",
        "dateCloture",
        "dateExercice",
        "dateDepot",
        "updatedAt",
        "createdAt",
    ]:
        v = meta.get(key)
        if v not in [None, ""]:
            candidates.append(v)

    for raw in candidates:
        s = str(raw).strip()
        if not s:
            continue
        m = re.search(r"(19|20)\d{2}", s)
        if m:
            return m.group(0)
    return ""


def make_doc_filename(
    famille: str,
    meta: Dict[str, Any],
    siren: str,
    denomination: str,
    forced_type_label: str = "",
) -> str:
    date_depot = valid_date(
        meta.get("dateDepot"),
        meta.get("updatedAt"),
        meta.get("createdAt"),
    ) or datetime.now().strftime("%Y-%m-%d")
    label = (forced_type_label or resolve_doc_type(meta, famille)["type_label"]).strip() or "Document"
    if famille.upper() != "ACTE" and label.upper() in {"NA", "N/A", "DOCUMENT"}:
        year_ref = extract_reference_year(meta)
        if year_ref:
            label = year_ref
    label = re.sub(r"\s+", " ", label).strip()[:90]
    denomination = sanitize_filename(denomination or "Entreprise")
    famille = "ACTE" if famille.upper() == "ACTE" else "COMPTES ANNUELS"
    base = f"{date_depot} - {siren} - {denomination} - {famille} - {label}"
    return sanitize_filename(base)


def make_attestation_filename(run_date: str, siren: str, denomination: str) -> str:
    return sanitize_filename(
        f"{run_date} - {siren} - {denomination or 'Entreprise'} - ATTESTATION RNE"
    )


def make_report_filename(run_date: str, siren: str, denomination: str) -> str:
    return sanitize_filename(
        f"{run_date} - {siren} - {denomination or 'Entreprise'} - Rapport Collecte RNE"
    )


def unique_zip_path(path: str, used_paths: set[str]) -> str:
    base, ext = os.path.splitext(path)
    candidate = path
    idx = 2
    while candidate in used_paths:
        candidate = f"{base} ({idx}){ext}"
        idx += 1
    used_paths.add(candidate)
    return candidate


def process_one_document(
    famille: str,
    meta: Dict[str, Any],
    pdf_bytes: bytes,
    siren: str,
    denomination: str,
) -> Dict[str, Any]:
    text = extract_pdf_text(pdf_bytes)
    descr = first_lines(text)
    dtype = resolve_doc_type(meta, famille)
    rdd_rows = extract_type_rdd_rows(meta) if famille.upper() == "ACTE" else []

    filename_base = make_doc_filename(famille, meta, siren, denomination, forced_type_label=dtype["type_label"])
    base = {
        "famille": famille,
        "filename_base": filename_base,
        "type_code": dtype["type_code"],
        "type_label": dtype["type_label"],
        "date_depot": valid_date(meta.get("dateDepot"), meta.get("updatedAt"), meta.get("createdAt")),
        "analysed": "OUI" if descr else "NON",
        "descriptif": descr,
        "doc_id": str(meta.get("id") or ""),
        "source": "ACTE" if famille.upper() == "ACTE" else "BILAN",
        "endpoint": "",
    }

    labels = [x.get("type_label", "").strip() for x in rdd_rows if str(x.get("type_label") or "").strip()]
    dedup_labels: List[str] = []
    seen_labels = set()
    for lbl in labels:
        k = lbl.casefold()
        if k in seen_labels:
            continue
        seen_labels.add(k)
        dedup_labels.append(lbl)
    nb_sub = max(1, len(dedup_labels))
    base["nb_actes_in_doc"] = nb_sub
    base["actes_labels"] = dedup_labels
    base["is_multi_actes"] = bool(famille.upper() == "ACTE" and nb_sub > 1)
    base["report_rows"] = [base]
    return base


def iter_nodes(root: Any) -> Iterable[Any]:
    yield root
    if isinstance(root, dict):
        for v in root.values():
            yield from iter_nodes(v)
    elif isinstance(root, list):
        for v in root:
            yield from iter_nodes(v)


def find_key_values(root: Any, target_key: str) -> List[Any]:
    out: List[Any] = []
    for node in iter_nodes(root):
        if isinstance(node, dict) and target_key in node and node[target_key] not in [None, ""]:
            out.append(node[target_key])
    return out


def first_value(root: Any, keys: List[str]) -> str:
    for k in keys:
        vals = find_key_values(root, k)
        if vals:
            v = vals[0]
            if isinstance(v, (dict, list)):
                continue
            return str(v)
    return ""


def deep_get(root: Dict[str, Any], path: str, default: Any = None) -> Any:
    cur: Any = root
    for key in path.split("."):
        if not isinstance(cur, dict):
            return default
        cur = cur.get(key)
        if cur is None:
            return default
    return cur


def non_empty(v: Any) -> bool:
    return v not in [None, "", [], {}]


def first_non_empty(*values: Any) -> Any:
    for v in values:
        if non_empty(v):
            return v
    return ""


def compact_dict(v: Any) -> Any:
    if isinstance(v, dict):
        out: Dict[str, Any] = {}
        for k, x in v.items():
            if k.endswith("Present"):
                continue
            cx = compact_dict(x)
            if non_empty(cx):
                out[k] = cx
        return out
    if isinstance(v, list):
        out_l = [compact_dict(x) for x in v]
        return [x for x in out_l if non_empty(x)]
    return v


def format_address(addr: Dict[str, Any]) -> str:
    if not isinstance(addr, dict):
        return ""
    parts = [
        str(addr.get("numVoie") or "").strip(),
        str(addr.get("typeVoie") or "").strip(),
        str(addr.get("voie") or "").strip(),
        str(addr.get("complementLocalisation") or "").strip(),
        str(addr.get("codePostal") or "").strip(),
        str(addr.get("commune") or "").strip(),
        str(addr.get("pays") or "").strip(),
    ]
    return " ".join([p for p in parts if p]).strip()


def extract_representants(content: Dict[str, Any]) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    pouvoirs = deep_get(content, "personneMorale.composition.pouvoirs", []) or []
    if not isinstance(pouvoirs, list):
        return out
    for p in pouvoirs:
        if not isinstance(p, dict):
            continue
        actif = bool(p.get("actif", False))
        if not actif:
            continue
        tpers = str(p.get("typeDePersonne") or "").strip().upper()
        role_code = str(p.get("roleEntreprise") or "").strip()
        role = resolve_label("roleEntreprise", role_code, fallback=role_code)
        row: Dict[str, Any] = {
            "type": tpers or "-",
            "role": role or "-",
            "actif": actif,
            "ville": "-",
            "code_postal": "-",
        }
        if tpers == "INDIVIDU":
            d = deep_get(p, "individu.descriptionPersonne", {}) or {}
            prenoms = d.get("prenoms") or []
            row["nom"] = normalize_text(d.get("nom") or "-")
            row["prenoms"] = " ".join([normalize_text(x) for x in prenoms if str(x or "").strip()]) or "-"
            row["denomination"] = "-"
            row["siren"] = "-"
            adr = deep_get(p, "individu.adresseDomicile", {}) or {}
            row["ville"] = str(adr.get("commune") or "-")
            row["code_postal"] = str(adr.get("codePostal") or "-")
        else:
            pm = first_non_empty(p.get("personneMorale"), p.get("entreprise"), {})
            if not isinstance(pm, dict):
                pm = {}
            row["nom"] = "-"
            row["prenoms"] = "-"
            row["denomination"] = normalize_text(first_non_empty(pm.get("denomination"), pm.get("raisonSociale"), "-"))
            row["siren"] = str(pm.get("siren") or "-")
            adr = first_non_empty(pm.get("adresse"), {}) if isinstance(pm, dict) else {}
            if isinstance(adr, dict):
                row["ville"] = str(adr.get("commune") or "-")
                row["code_postal"] = str(adr.get("codePostal") or "-")
        out.append(row)
        if len(out) >= 15:
            break
    return out


def extract_etablissements(content: Dict[str, Any]) -> Dict[str, Any]:
    principal = deep_get(content, "personneMorale.etablissementPrincipal", {}) or {}
    pdesc = (principal.get("descriptionEtablissement") or {}) if isinstance(principal, dict) else {}
    padr = (principal.get("adresse") or {}) if isinstance(principal, dict) else {}
    pacts = (principal.get("activites") or []) if isinstance(principal, dict) else []
    pape = ""
    if isinstance(pacts, list):
        for a in pacts:
            if isinstance(a, dict) and a.get("indicateurPrincipal"):
                pape = str(a.get("codeApe") or "")
                break
    principal_out = {
        "siret": str(pdesc.get("siret") or ""),
        "commune": str(padr.get("commune") or ""),
        "code_postal": str(padr.get("codePostal") or ""),
        "statut": str(pdesc.get("statutPourFormalite") or ""),
        "ape": pape,
        "principal": True,
        "date_fermeture": "",
    }
    synth = [principal_out] if principal_out["siret"] else []
    return {"principal": principal_out, "liste": synth[:10]}


def extract_observations_rcs(content: Dict[str, Any]) -> List[Dict[str, str]]:
    obs = deep_get(content, "personneMorale.observations.rcs", []) or []
    out: List[Dict[str, str]] = []
    if not isinstance(obs, list):
        return out
    for item in obs:
        if not isinstance(item, dict):
            continue
        texte = str(first_non_empty(item.get("texte"), item.get("observation"), "")).strip()
        if not texte:
            continue
        out.append(
            {
                "date": valid_date(item.get("dateAjout"), item.get("dateGreffe"), item.get("dateObservation")),
                "texte": texte,
            }
        )
        if len(out) >= 10:
            break
    return out


def parse_company_rne(company_json: dict) -> dict:
    payload = company_json.get("data", company_json) if isinstance(company_json, dict) else {}
    formality = payload.get("formality", {}) if isinstance(payload, dict) else {}
    content = formality.get("content", {}) if isinstance(formality, dict) else {}
    entreprise = deep_get(content, "personneMorale.identite.entreprise", {}) or {}
    desc = deep_get(content, "personneMorale.identite.description", {}) or {}
    pub = deep_get(content, "personneMorale.identite.publicationLegale", {}) or {}
    adr_ent = deep_get(content, "personneMorale.adresseEntreprise.adresse", {}) or {}
    carac = deep_get(content, "personneMorale.adresseEntreprise.caracteristiques", {}) or {}
    etab_principal = deep_get(content, "personneMorale.etablissementPrincipal", {}) or {}
    etab_pr_desc = (etab_principal.get("descriptionEtablissement") or {}) if isinstance(etab_principal, dict) else {}
    acts = (etab_principal.get("activites") or []) if isinstance(etab_principal, dict) else []
    powers = deep_get(content, "personneMorale.composition.pouvoirs", []) or []
    historique = formality.get("historique", []) if isinstance(formality, dict) else []
    obs_rcs_raw = deep_get(content, "personneMorale.observations.rcs", []) or []

    activites: List[Dict[str, Any]] = []
    if isinstance(acts, list):
        for a in acts:
            if not isinstance(a, dict):
                continue
            activites.append(
                {
                    "indicateurPrincipal": a.get("indicateurPrincipal"),
                    "dateDebut": valid_date(a.get("dateDebut")),
                    "codeApe": a.get("codeApe"),
                    "descriptionDetaillee": normalize_text(a.get("descriptionDetaillee")),
                    "precisionAutre": normalize_text(a.get("precisionAutre")),
                    "formeExercice": str(a.get("formeExercice") or ""),
                    "formeExerciceLibelle": resolve_label("formeExercice", a.get("formeExercice")),
                    "exerciceActivite": str(a.get("exerciceActivite") or ""),
                    "exerciceActiviteLibelle": resolve_label("exerciceActivite", a.get("exerciceActivite")),
                    "typeOrigine": str(deep_get(a, "origine.typeOrigine") or ""),
                    "typeOrigineLibelle": resolve_label("typeOrigine", deep_get(a, "origine.typeOrigine")),
                }
            )
            if len(activites) >= 3:
                break

    reps: List[Dict[str, Any]] = []
    if isinstance(powers, list):
        for p in powers:
            if not isinstance(p, dict):
                continue
            if not bool(p.get("actif", False)):
                continue
            tpers = str(p.get("typeDePersonne") or "").strip().upper()
            role_code_rne = str(p.get("roleEntreprise") or "")
            row: Dict[str, Any] = {
                "roleEntreprise": role_code_rne,
                "roleEntrepriseLibelle": resolve_label("roleEntreprise", role_code_rne, fallback=role_code_rne),
                "typeDePersonne": tpers,
                "typeDePersonneLibelle": resolve_label("typeDePersonne", tpers, fallback=tpers),
                "actif": bool(p.get("actif", False)),
                "nom": "",
                "prenoms": "",
                "dateDeNaissance": "",
                "commune": "",
                "codePostal": "",
            }
            if tpers == "INDIVIDU":
                d = deep_get(p, "individu.descriptionPersonne", {}) or {}
                adr = deep_get(p, "individu.adresseDomicile", {}) or {}
                prenoms = d.get("prenoms") or []
                row["nom"] = normalize_text(d.get("nom") or "")
                row["prenoms"] = " ".join([normalize_text(x) for x in prenoms if str(x or "").strip()]) or ""
                row["dateDeNaissance"] = str(d.get("dateDeNaissance") or "")
                row["commune"] = str(adr.get("commune") or "")
                row["codePostal"] = str(adr.get("codePostal") or "")
            else:
                pm = first_non_empty(p.get("personneMorale"), p.get("entreprise"), {})
                if not isinstance(pm, dict):
                    pm = {}
                den = normalize_text(first_non_empty(pm.get("denomination"), pm.get("raisonSociale"), pm.get("nom"), ""))
                psiren = str(pm.get("siren") or "")
                # For legal entities, expose company info in table columns without changing schema.
                row["nom"] = den
                row["prenoms"] = f"SIREN {psiren}" if psiren else ""
                adr_pm = first_non_empty(pm.get("adresse"), pm.get("adresseEntreprise"), {})
                if isinstance(adr_pm, dict):
                    row["commune"] = str(first_non_empty(adr_pm.get("commune"), ""))
                    row["codePostal"] = str(first_non_empty(adr_pm.get("codePostal"), ""))
            reps.append(row)
            if len(reps) >= 10:
                break

    hist_out: List[Dict[str, Any]] = []
    if isinstance(historique, list):
        for h in historique:
            if not isinstance(h, dict):
                continue
            hist_out.append(
                {
                    "dateIntegration": valid_date(h.get("dateIntegration")),
                    "codeEvenement": str(h.get("codeEvenement") or ""),
                    "numeroLiasse": str(h.get("numeroLiasse") or ""),
                    "dateEffet": valid_date(h.get("dateEffet")),
                }
            )
            if len(hist_out) >= 5:
                break

    situation_juridique = {
        "statutRcs": str(first_value(content, ["statutRcs", "statutRegistre"]) or ""),
        "dateCessationActivite": valid_date(
            first_value(
                content,
                ["dateCessationTotaleActivite", "dateCessationActivite", "dateCessation"],
            )
        ),
        "dateDissolution": valid_date(first_value(content, ["dateDissolution"])),
        "dateRadiation": valid_date(first_value(content, ["dateRadiation"])),
        "dateClotureLiquidation": valid_date(
            first_value(content, ["dateClotureLiquidation", "dateCloture"])
        ),
        "indicateurCessationActivite": first_value(
            content, ["indicateurCessationActivite", "cessationActivite"]
        ),
        "indicateurRadiation": first_value(
            content, ["indicateurRadiationRcs", "entrepriseRadiee"]
        ),
    }

    observations_rcs: List[Dict[str, str]] = []
    if isinstance(obs_rcs_raw, list):
        for o in obs_rcs_raw:
            if not isinstance(o, dict):
                continue
            txt = normalize_text(
                first_non_empty(o.get("texte"), o.get("observation"), "")
            ).strip()
            if not txt:
                continue
            observations_rcs.append(
                {
                    "date": valid_date(
                        o.get("dateAjout"), o.get("dateGreffe"), o.get("dateObservation")
                    ),
                    "texte": txt,
                }
            )
            if len(observations_rcs) >= 10:
                break

    parsed = {
        "meta": {
            "updatedAt": valid_date(payload.get("updatedAt")),
            "formality_created": valid_date(formality.get("created")),
            "formality_updated": valid_date(formality.get("updated")),
            "date_creation": valid_date(
                deep_get(content, "natureCreation.dateCreation"),
                first_value(content, ["dateCreation", "dateCreationEntreprise"]),
            ),
            "date_immatriculation": valid_date(
                first_value(content, ["dateImmatriculation", "dateImmatriculationRcs"])
            ),
        },
        "identite": {
            "siren": str(entreprise.get("siren") or formality.get("siren") or payload.get("siren") or ""),
            "denomination": normalize_text(entreprise.get("denomination") or ""),
            "sigle": normalize_text(entreprise.get("sigle") or ""),
            "formeJuridique": str(entreprise.get("formeJuridique") or formality.get("formeJuridique") or ""),
            "formeJuridiqueLibelle": resolve_label(
                "formeJuridique",
                entreprise.get("formeJuridique") or formality.get("formeJuridique"),
            ),
            "nicSiege": str(entreprise.get("nicSiege") or ""),
            "codeApe": str(entreprise.get("codeApe") or ""),
        },
        "statut_diffusion": {
            "diffusionINSEE": str(formality.get("diffusionINSEE") or ""),
            "diffusionCommerciale": formality.get("diffusionCommerciale"),
            "typePersonne": str(formality.get("typePersonne") or ""),
            "typePersonneLibelle": resolve_person_type_label(formality.get("typePersonne")),
        },
        "parametres_juridiques": {
            "duree": desc.get("duree"),
            "datePremiereCloture": valid_date(desc.get("datePremiereCloture")),
            "montantCapital": desc.get("montantCapital"),
            "deviseCapital": desc.get("deviseCapital"),
            "capitalVariable": desc.get("capitalVariable"),
            "societeMission": desc.get("societeMission"),
            "indicateurAssocieUnique": desc.get("indicateurAssocieUnique"),
            "depotDemandeAcre": desc.get("depotDemandeAcre"),
        },
        "objet_social": normalize_text(desc.get("objet") or ""),
        "publication_legale": {
            "datePublication": valid_date(pub.get("datePublication")),
            "journalPublication": normalize_text(pub.get("journalPublication") or ""),
            "supportPublication": normalize_text(first_non_empty(pub.get("support"), pub.get("supportPublication"), "")),
            "numeroPublication": normalize_text(first_non_empty(pub.get("numeroPublication"), pub.get("numero"), "")),
        },
        "domiciliation": {
            "domiciliataire": bool(carac.get("domiciliataire", False)),
            "domiciliation_siren": str(deep_get(content, "personneMorale.adresseEntreprise.entrepriseDomiciliataire.siren") or ""),
            "domiciliation_denomination": normalize_text(
                deep_get(content, "personneMorale.adresseEntreprise.entrepriseDomiciliataire.denomination") or ""
            ),
        },
        "siege": {
            "numVoie": str(adr_ent.get("numVoie") or ""),
            "typeVoie": str(adr_ent.get("typeVoie") or ""),
            "typeVoieLibelle": resolve_label("typeVoie", adr_ent.get("typeVoie")),
            "voie": str(adr_ent.get("voie") or ""),
            "complementLocalisation": str(adr_ent.get("complementLocalisation") or ""),
            "codePostal": str(adr_ent.get("codePostal") or ""),
            "commune": str(adr_ent.get("commune") or ""),
            "codeInseeCommune": str(adr_ent.get("codeInseeCommune") or ""),
            "pays": str(adr_ent.get("pays") or adr_ent.get("codePays") or ""),
            "paysLibelle": resolve_label("pays", adr_ent.get("codePays"), fallback=str(adr_ent.get("pays") or "")),
            "indicateurValidationBANPresent": adr_ent.get("indicateurValidationBANPresent"),
        },
        "etablissement_principal": {
            "siret": str(etab_pr_desc.get("siret") or ""),
            "statutPourFormalite": str(etab_pr_desc.get("statutPourFormalite") or ""),
            "indicateurEtablissementPrincipal": etab_pr_desc.get("indicateurEtablissementPrincipal"),
        },
        "activites": activites,
        "representants_actifs": reps,
        "historique": hist_out,
        "situation_juridique": situation_juridique,
        "observations_rcs": observations_rcs,
    }
    return compact_dict(parsed)


def build_company_summary(
    company_json: Dict[str, Any],
    siren: str,
    bilans: List[Dict[str, Any]],
    actes: List[Dict[str, Any]],
) -> Dict[str, Any]:
    parsed = parse_company_rne(company_json)
    identite = parsed.get("identite", {}) if isinstance(parsed, dict) else {}
    domic = parsed.get("domiciliation", {}) if isinstance(parsed, dict) else {}
    situation = parsed.get("situation_juridique", {}) if isinstance(parsed, dict) else {}
    observations = parsed.get("observations_rcs", []) if isinstance(parsed, dict) else []
    summary: Dict[str, Any] = {
        "denomination": str(identite.get("denomination") or ""),
        "siren": str(identite.get("siren") or siren),
        "nb_bilans": len(bilans),
        "nb_comptes_annuels": len(bilans),
        "nb_documents_juridiques": len(actes),
        "nb_actes": len(actes),
        "nb_docs": len(bilans) + len(actes),
        "points_factuels": [],
        "parsed": parsed,
    }
    if len(bilans) == 0:
        summary["points_factuels"].append("Aucun bilan disponible")
    if len(actes) == 0:
        summary["points_factuels"].append("Aucun acte disponible")
    if bool(domic.get("domiciliataire", False)):
        summary["points_factuels"].append("Domiciliataire: OUI")
    if str(situation.get("dateRadiation") or "").strip():
        summary["points_factuels"].append(
            f"Radiation enregistrée le {situation.get('dateRadiation')}"
        )
    if str(situation.get("dateClotureLiquidation") or "").strip():
        summary["points_factuels"].append(
            f"Clôture de liquidation le {situation.get('dateClotureLiquidation')}"
        )
    if str(situation.get("dateDissolution") or "").strip():
        summary["points_factuels"].append(
            f"Dissolution enregistrée le {situation.get('dateDissolution')}"
        )
    if str(situation.get("dateCessationActivite") or "").strip():
        summary["points_factuels"].append(
            f"Cessation d'activité le {situation.get('dateCessationActivite')}"
        )
    for o in observations:
        txt = str(o.get("texte") or "")
        if "liquidation" in txt.lower():
            summary["points_factuels"].append(
                f"Observation RCS mentionne une liquidation ({o.get('date') or '-'})"
            )
            break
    return compact_dict(summary)


# ══════════════════════════════════════════════════════════════════════════════
# PALETTE — RAPPORT RNE
# ══════════════════════════════════════════════════════════════════════════════
_C_ORANGE   = HexColor("#EF8829")
_C_TEAL     = HexColor("#017e84")
_C_NAVY     = HexColor("#1a2744")
_C_NAVY_L   = HexColor("#243456")
_C_TXT      = HexColor("#2c2c2c")
_C_MUTED    = HexColor("#7a7470")
_C_BORDER   = HexColor("#e8e4df")
_C_LGRAY    = HexColor("#f3f1ee")
_C_CARD     = HexColor("#ffffff")
_C_RED      = HexColor("#c0392b")
_C_GREEN    = HexColor("#27ae60")
_C_ALERT_BG = HexColor("#fff8f0")
_C_BG       = HexColor("#FAF8F6")
_PW, _PH    = landscape(A4)

# DOCX colours
_D_ORANGE = RGBColor(0xEF, 0x88, 0x29)
_D_TEAL   = RGBColor(0x01, 0x7E, 0x84)
_D_NAVY   = RGBColor(0x1A, 0x27, 0x44)
_D_MUTED  = RGBColor(0x7A, 0x74, 0x70)
_D_GREEN  = RGBColor(0x27, 0xAE, 0x60)
_D_RED    = RGBColor(0xC0, 0x39, 0x2B)
_D_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)


# ── PDF — callbacks de page ────────────────────────────────────────────────
def _rne_on_cover(c: Any, doc: Any) -> None:
    c.saveState()
    c.setFillColor(_C_NAVY); c.rect(0, 0, _PW, _PH, fill=1, stroke=0)
    c.setFillColor(_C_ORANGE); c.rect(0, _PH - 7 * mm, _PW, 7 * mm, fill=1, stroke=0)
    c.setStrokeColor(_C_ORANGE); c.setLineWidth(1.5)
    c.line(2.5 * cm, 7 * cm, _PW - 2.5 * cm, 7 * cm)
    c.restoreState()


def _rne_on_closing(c: Any, doc: Any) -> None:
    c.saveState()
    c.setFillColor(_C_NAVY); c.rect(0, 0, _PW, _PH, fill=1, stroke=0)
    c.setFillColor(_C_ORANGE); c.rect(0, 0, _PW, 5 * mm, fill=1, stroke=0)
    c.restoreState()


def _rne_on_content(c: Any, doc: Any) -> None:
    c.saveState()
    c.setFillColor(_C_BG); c.rect(0, 0, _PW, _PH, fill=1, stroke=0)
    c.setStrokeColor(_C_ORANGE); c.setLineWidth(1)
    c.line(1.5 * cm, _PH - 1.0 * cm, _PW - 1.5 * cm, _PH - 1.0 * cm)
    c.setFont("Helvetica", 6.5); c.setFillColor(_C_MUTED)
    c.drawString(1.5 * cm, _PH - 0.85 * cm, "Gidoo \u2014 Rapport de collecte RNE")
    dn = getattr(doc, "_dn", ""); sr = getattr(doc, "_sr", "")
    c.drawRightString(_PW - 1.5 * cm, _PH - 0.85 * cm, f"{dn} \u2014 SIREN {sr}")
    c.setStrokeColor(_C_BORDER); c.setLineWidth(0.4)
    c.line(1.5 * cm, 1.2 * cm, _PW - 1.5 * cm, 1.2 * cm)
    c.setFont("Helvetica", 6.5); c.setFillColor(_C_MUTED)
    rd = getattr(doc, "_rd", "")
    c.drawString(1.5 * cm, 0.8 * cm,
                 f"Rapport du {rd} \u2014 Source : INPI / Registre National des Entreprises")
    c.drawRightString(_PW - 1.5 * cm, 0.8 * cm, f"Page {c.getPageNumber()}")
    c.restoreState()


# ── PDF — AlertBox ─────────────────────────────────────────────────────────
class _AlertBox(Flowable):
    def __init__(self, width: float, paras: List[Any], title: str = "Points d\u2019attention") -> None:
        Flowable.__init__(self)
        self.bw, self.paras, self.title, self._pad = width, paras, title, 12
        aw = width - 2 * self._pad
        self._hs = [p.wrap(aw, 10000)[1] for p in paras]
        self._th = 26
        self.height = self._th + sum(self._hs) + 2 * self._pad + len(paras) * 3
        self.width = width

    def draw(self) -> None:
        c = self.canv; p, r = self._pad, 5; w, h = self.bw, self.height
        c.setFillColor(_C_ALERT_BG); c.setStrokeColor(_C_ORANGE); c.setLineWidth(0.5)
        c.roundRect(0, 0, w, h, r, fill=1, stroke=1)
        pth = c.beginPath()
        pth.moveTo(0, h - self._th); pth.lineTo(0, h - r)
        pth.arcTo(0, h - 2 * r, 2 * r, h, startAng=90, extent=90)
        pth.lineTo(w - r, h)
        pth.arcTo(w - 2 * r, h - 2 * r, w, h, startAng=0, extent=90)
        pth.lineTo(w, h - self._th); pth.close()
        c.setFillColor(_C_ORANGE); c.drawPath(pth, fill=1, stroke=0)
        c.setFillColor(white); c.setFont("Helvetica-Bold", 9.5)
        c.drawString(p, h - self._th + 7, self.title)
        y = h - self._th - p; aw = w - 2 * p
        for para in self.paras:
            _, ph = para.wrap(aw, 10000); para.drawOn(c, p, y - ph); y -= ph + 3


# ── PDF — styles ───────────────────────────────────────────────────────────
def _rne_styles() -> Any:
    s = getSampleStyleSheet(); a = s.add
    a(ParagraphStyle("RCvT", fontName="Helvetica-Bold", fontSize=34, leading=42, textColor=white, alignment=TA_LEFT))
    a(ParagraphStyle("RCvS", fontName="Helvetica", fontSize=15, leading=22, textColor=HexColor("#b0c4de"), alignment=TA_LEFT))
    a(ParagraphStyle("RCvM", fontName="Helvetica", fontSize=11, leading=16, textColor=HexColor("#8899bb"), alignment=TA_LEFT))
    a(ParagraphStyle("RTocT", fontName="Helvetica-Bold", fontSize=22, leading=28, textColor=_C_NAVY, spaceAfter=16))
    a(ParagraphStyle("RToc1", fontName="Helvetica-Bold", fontSize=11, leading=20, textColor=_C_TXT))
    a(ParagraphStyle("RToc2", fontName="Helvetica", fontSize=10, leading=18, textColor=_C_MUTED, leftIndent=18))
    a(ParagraphStyle("RH1", fontName="Helvetica-Bold", fontSize=18, leading=24, textColor=_C_NAVY, spaceAfter=8))
    a(ParagraphStyle("RH2", fontName="Helvetica-Bold", fontSize=13, leading=18, textColor=_C_TEAL, spaceBefore=12, spaceAfter=6))
    a(ParagraphStyle("RBd", fontName="Helvetica", fontSize=9.5, leading=13.5, textColor=_C_TXT, alignment=TA_JUSTIFY, spaceAfter=4))
    a(ParagraphStyle("RBdS", fontName="Helvetica", fontSize=8.5, leading=12, textColor=_C_TXT, spaceAfter=2))
    a(ParagraphStyle("RLbl", fontName="Helvetica-Bold", fontSize=9, leading=12, textColor=_C_MUTED))
    a(ParagraphStyle("RVal", fontName="Helvetica", fontSize=9.5, leading=13, textColor=_C_TXT))
    a(ParagraphStyle("RTh", fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=white))
    a(ParagraphStyle("RTd", fontName="Helvetica", fontSize=8.5, leading=12, textColor=_C_TXT))
    a(ParagraphStyle("RBgO", fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=_C_GREEN))
    a(ParagraphStyle("RBgN", fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=_C_RED))
    a(ParagraphStyle("RAlt", fontName="Helvetica", fontSize=9.5, leading=14, textColor=_C_TXT, leftIndent=14))
    a(ParagraphStyle("RDocT", fontName="Helvetica-Bold", fontSize=11, leading=15, textColor=_C_NAVY, spaceBefore=8, spaceAfter=4))
    a(ParagraphStyle("RClT", fontName="Helvetica-Bold", fontSize=16, leading=22, textColor=white, alignment=TA_LEFT))
    a(ParagraphStyle("RClS", fontName="Helvetica", fontSize=10, leading=14, textColor=white, alignment=TA_LEFT))
    return s


def _rx(t: Any) -> str:
    s = str(t) if t else ""
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _rne_kv_table(rows: List[Any], sty: Any, col1: float = 5.5 * cm, total_w: Optional[float] = None) -> Table:
    if total_w is None:
        total_w = _PW - 3.5 * cm
    col2 = total_w - col1
    data = [[RLPara(f"<b>{_rx(lbl)}</b>", sty["RLbl"]), RLPara(_rx(str(val)), sty["RVal"])]
            for lbl, val in rows]
    ts: List[Any] = [
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 10), ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("LINEBELOW", (0, 0), (-1, -2), 0.3, _C_BORDER),
    ]
    for i in range(len(data)):
        ts.append(("BACKGROUND", (0, i), (-1, i), _C_LGRAY if i % 2 == 0 else _C_CARD))
    t = Table(data, colWidths=[col1, col2])
    t.setStyle(TableStyle(ts))
    return t


# ── Conversion company_summary → rne_data ─────────────────────────────────
def _to_rne_data(company_summary: Dict[str, Any], siren: str, denomination: str) -> Dict[str, Any]:
    parsed = (company_summary.get("parsed") or {})
    ident  = parsed.get("identite", {}) or {}
    pj     = parsed.get("parametres_juridiques", {}) or {}
    siege  = parsed.get("siege", {}) or {}
    etp    = parsed.get("etablissement_principal", {}) or {}
    acts   = parsed.get("activites", []) or []
    reps   = parsed.get("representants_actifs", []) or []
    meta   = parsed.get("meta", {}) or {}
    hist   = parsed.get("historique", []) or []
    pub    = parsed.get("publication_legale", {}) or {}
    statut = parsed.get("statut_diffusion", {}) or {}

    def _push(rows: List[Any], label: str, value: Any) -> None:
        sval = normalize_text(value) if isinstance(value, str) else str(value or "").strip()
        if sval:
            rows.append((label, sval))

    def _bool_or_empty(value: Any) -> str:
        if isinstance(value, bool):
            return "Oui" if value else "Non"
        return ""

    def _format_capital_eur(value: Any) -> str:
        if value in [None, ""]:
            return ""
        sval = str(value).strip().replace(" ", "").replace(",", ".")
        try:
            amount = float(sval)
        except Exception:
            raw = str(value).strip()
            return f"{raw} €" if raw else ""
        if abs(amount - round(amount)) < 1e-9:
            int_part = int(round(amount))
            return f"{format(int_part, ',').replace(',', '.')} €"
        int_part = int(amount)
        dec_part = int(round((amount - int_part) * 100))
        return f"{format(int_part, ',').replace(',', '.')},{dec_part:02d} €"

    def _format_duree_ans(value: Any) -> str:
        s = str(value or "").strip()
        if not s:
            return ""
        return s if "ans" in s.lower() else f"{s} ans"

    def _role_sort_key(role_lbl: str) -> int:
        r = (role_lbl or "").casefold()
        if "président" in r or "president" in r:
            return 0
        if "directeur général" in r or "directeur general" in r:
            return 1
        return 50

    def _cac_sort_key(role_lbl: str) -> int:
        r = (role_lbl or "").casefold()
        if "titulaire" in r:
            return 0
        if "suppléant" in r or "suppleant" in r:
            return 1
        return 50

    sections: List[Dict[str, Any]] = []
    section_order: List[Any] = []

    den = normalize_text(ident.get("denomination") or denomination or "")

    # A.1 Identification
    a1: List[Any] = []
    _push(a1, "Dénomination", den)
    _push(a1, "Sigle", ident.get("sigle"))
    _push(a1, "SIREN", siren)
    _push(a1, "SIRET du siège", etp.get("siret"))
    forme_lbl = (ident.get("formeJuridiqueLibelle") or ident.get("formeJuridique") or "").strip()
    _push(a1, "Forme juridique", forme_lbl)
    type_pers = str(
        statut.get("typePersonneLibelle")
        or resolve_person_type_label(statut.get("typePersonne"))
        or statut.get("typePersonne")
        or ""
    ).strip()
    _push(a1, "Type de personne", type_pers)
    if a1:
        sections.append({"id": "A1_identification", "title": "A.1 Identification de l'entreprise", "rows": a1})
        section_order.append(("A.1", "Identification de l'entreprise"))

    # A.2 Activité
    a2: List[Any] = []
    act_principale = next((a for a in acts if a.get("indicateurPrincipal")), acts[0] if acts else {})
    ape_code = str(ident.get("codeApe") or act_principale.get("codeApe") or "").strip()
    ape_lib = normalize_text(act_principale.get("descriptionDetaillee") or "")
    seen_activity_values: set[str] = set()

    def _activity_key(value: str) -> str:
        return re.sub(r"\s+", " ", str(value or "").strip()).casefold()

    if ape_code:
        _push(a2, "Code APE", ape_code)
    if ape_lib:
        _push(a2, "Libellé APE", ape_lib)
        seen_activity_values.add(_activity_key(ape_lib))
    objet = normalize_text(parsed.get("objet_social") or "").strip()
    if objet:
        _push(a2, "Activité principale", objet[:600])
        seen_activity_values.add(_activity_key(objet[:600]))
    acts_sec = [a for a in acts if not a.get("indicateurPrincipal")]
    for i, a in enumerate(acts_sec[:3], start=1):
        desc_sec = normalize_text(a.get("descriptionDetaillee") or a.get("precisionAutre") or "").strip()
        if desc_sec:
            code_sec = str(a.get("codeApe") or "").strip()
            label_sec = f"{code_sec} - {desc_sec}" if code_sec else desc_sec
            # Avoid duplicates with Libellé APE / Activité principale or already listed secondary entries.
            dedup_target = desc_sec if code_sec and label_sec.startswith(f"{code_sec} - ") else label_sec
            k = _activity_key(dedup_target)
            if k in seen_activity_values:
                continue
            seen_activity_values.add(k)
            _push(a2, f"Activité secondaire {i}", label_sec)
    if act_principale:
        _push(a2, "Forme d'exercice", act_principale.get("formeExerciceLibelle") or act_principale.get("formeExercice"))
        _push(a2, "Origine de l'activité", act_principale.get("typeOrigineLibelle") or act_principale.get("typeOrigine"))
    if a2:
        sections.append({"id": "A2_activite", "title": "A.2 Activité", "rows": a2})
        section_order.append(("A.2", "Activité"))

    # A.3 Siège et coordonnées
    a3: List[Any] = []
    adresse_complete = " ".join(
        [x for x in [str(siege.get("numVoie") or "").strip(),
                     str(siege.get("typeVoieLibelle") or siege.get("typeVoie") or "").strip(),
                     str(siege.get("voie") or "").strip()] if x]
    ).strip()
    _push(a3, "Adresse complète", adresse_complete)
    _push(a3, "Complément d'adresse", siege.get("complementLocalisation"))
    _push(a3, "Code postal", siege.get("codePostal"))
    _push(a3, "Commune", siege.get("commune"))
    _push(a3, "Pays", siege.get("paysLibelle") or siege.get("pays"))
    ban_value = _bool_or_empty(siege.get("indicateurValidationBANPresent"))
    _push(a3, "Validation BAN", ban_value)
    if a3:
        sections.append({"id": "A3_siege", "title": "A.3 Siège et coordonnées", "rows": a3})
        section_order.append(("A.3", "Siège et coordonnées"))

    # A.4 Caractéristiques juridiques et financières
    a4: List[Any] = []
    if pj.get("montantCapital") is not None:
        cap_str = _format_capital_eur(pj.get("montantCapital"))
        _push(a4, "Capital social", cap_str)
    if pj.get("duree") is not None:
        _push(a4, "Durée", _format_duree_ans(pj.get("duree")))
    _push(a4, "Date de clôture", format_date_fr_optional(pj.get("datePremiereCloture")))
    _push(a4, "Associé unique", _bool_or_empty(pj.get("indicateurAssocieUnique")))
    _push(a4, "Capital variable", _bool_or_empty(pj.get("capitalVariable")))
    _push(a4, "Société à mission", _bool_or_empty(pj.get("societeMission")))
    if a4:
        sections.append({"id": "A4_juridique", "title": "A.4 Caractéristiques juridiques et financières", "rows": a4})
        section_order.append(("A.4", "Caractéristiques juridiques et financières"))

    # A.5 Représentants (hors commissaires aux comptes)
    a5: List[Any] = []
    a5_cac: List[Any] = []
    reps_pairs: List[Any] = []
    for idx, rep in enumerate(reps[:20]):
        role_lbl = (rep.get("roleEntrepriseLibelle") or rep.get("roleEntreprise") or "Rôle").strip()
        role_low = role_lbl.casefold()
        tpers = str(rep.get("typeDePersonne") or "").strip().upper()
        type_lbl = "Personne morale" if tpers != "INDIVIDU" else "Personne physique"

        details: List[str] = [type_lbl]
        if tpers == "INDIVIDU":
            nom_complet = " ".join(
                [x for x in [str(rep.get("nom") or "").strip(), str(rep.get("prenoms") or "").strip()] if x]
            ).strip()
            if nom_complet:
                details.append(nom_complet)
            ddn = format_birth_month_year(rep.get("dateDeNaissance"))
            if ddn and ddn != "-":
                details.append(f"Né(e) : {ddn}")
        else:
            den = str(rep.get("nom") or "").strip()
            siren_rep = str(rep.get("prenoms") or "").strip()
            if den and siren_rep:
                details.append(f"{den} | {siren_rep}")
            elif den:
                details.append(den)
            elif siren_rep:
                details.append(siren_rep)

        if rep.get("commune"):
            details.append(f"Commune : {rep.get('commune')}")
        if rep.get("codePostal"):
            details.append(f"Code postal : {rep.get('codePostal')}")
        if not details:
            continue
        reps_pairs.append((idx, role_lbl, " | ".join(details), role_low))

    general_rows = [
        (role_lbl, detail, idx)
        for idx, role_lbl, detail, role_low in reps_pairs
        if "commissaire aux comptes" not in role_low
    ]
    general_rows.sort(key=lambda x: (_role_sort_key(x[0]), x[2]))
    for role_lbl, detail, _idx in general_rows:
        a5.append((role_lbl, detail))

    cac_rows = [
        (role_lbl, detail, idx)
        for idx, role_lbl, detail, role_low in reps_pairs
        if "commissaire aux comptes" in role_low
    ]
    cac_rows.sort(key=lambda x: (_cac_sort_key(x[0]), x[2]))
    for role_lbl, detail, _idx in cac_rows:
        a5_cac.append((role_lbl, detail))

    if a5:
        sections.append({"id": "A5_representants", "title": "A.5 Représentants", "rows": a5})
        section_order.append(("A.5", "Représentants"))

    # Numérotation lisible sans sous-sections :
    # - si CAC présents : section "A.6 Commissaires aux comptes"
    # - publication passe en section suivante
    # - si pas de CAC : publication prend la section suivante disponible
    pub_section_num = "A.6" if a5 else "A.5"
    if a5_cac:
        cac_section_num = "A.6" if a5 else "A.5"
        sections.append({"id": "A5_cac", "title": f"{cac_section_num} Commissaires aux comptes", "rows": a5_cac})
        section_order.append((cac_section_num, "Commissaires aux comptes"))
        pub_section_num = "A.7" if a5 else "A.6"

    # A.6 Publication / historique utile
    a6: List[Any] = []
    creation = meta.get("date_creation") or meta.get("formality_created")
    _push(a6, "Date de création", format_date_fr_optional(creation))
    immat = meta.get("date_immatriculation")
    if not immat:
        for h in hist:
            if h.get("dateEffet"):
                immat = h.get("dateEffet")
                break
    if immat and immat != creation:
        _push(a6, "Date d'immatriculation", format_date_fr_optional(immat))
    _push(a6, "Date de dernière mise à jour", format_date_fr_optional(meta.get("updatedAt")))
    _push(a6, "Journal d'annonces légales", pub.get("journalPublication"))
    _push(a6, "Date de publication", format_date_fr_optional(pub.get("datePublication")))
    _push(a6, "Support de publication", pub.get("supportPublication"))
    _push(a6, "Référence de publication", pub.get("numeroPublication"))
    if a6:
        sections.append({"id": "A6_publication", "title": f"{pub_section_num} Publication / historique utile", "rows": a6})
        section_order.append((pub_section_num, "Publication / historique utile"))

    return {
        "denomination": den or denomination,
        "siren": siren,
        "points_attention": list(company_summary.get("points_factuels") or []),
        "sections": sections,
        "section_order": section_order,
    }


def _norm_doc_results(doc_results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out = []
    for d in doc_results:
        nb_actes = int(d.get("nb_actes_in_doc") or 1)
        labels = d.get("actes_labels") if isinstance(d.get("actes_labels"), list) else []
        out.append({
            "famille":          d.get("famille") or d.get("source") or "",
            "filename_base":    d.get("filename_base", "Document"),
            "typeDocument":     d.get("type_code") or "",
            "typeLibelle":      d.get("type_label") or "",
            "date_depot":       d.get("date_depot") or "",
            "document_analyse": d.get("analysed") or ("OUI" if d.get("descriptif") else "NON"),
            "descriptif":       d.get("descriptif") or "",
            "nb_actes_in_doc":  nb_actes,
            "actes_labels":     labels,
            "is_multi_actes":   bool(d.get("is_multi_actes") or nb_actes > 1),
        })
    return out


# ══════════════════════════════════════════════════════════════════════════════
# GÉNÉRATION PDF — template Gidoo
# ══════════════════════════════════════════════════════════════════════════════
def generate_report_pdf(
    siren: str,
    denomination: str,
    company_summary: Dict[str, Any],
    doc_results: List[Dict[str, Any]],
    run_date: str,
) -> bytes:
    rne    = _to_rne_data(company_summary, siren, denomination)
    docs   = _norm_doc_results(doc_results)
    buf    = io.BytesIO()
    sty    = _rne_styles()
    cw     = _PW - 3.0 * cm

    doc = BaseDocTemplate(buf, pagesize=landscape(A4),
                          topMargin=1.5 * cm, bottomMargin=1.5 * cm,
                          leftMargin=1.5 * cm, rightMargin=1.5 * cm)
    doc._dn = denomination; doc._sr = siren; doc._rd = run_date  # type: ignore[attr-defined]

    doc.addPageTemplates([
        PageTemplate("cover",   frames=[Frame(2.5*cm, 2*cm, _PW-5*cm, _PH-4*cm, id="cf")], onPage=_rne_on_cover),
        PageTemplate("content", frames=[Frame(1.5*cm, 1.5*cm, cw, _PH-3.2*cm, id="ct")],  onPage=_rne_on_content),
        PageTemplate("closing", frames=[Frame(2*cm, 2*cm, _PW-4*cm, _PH-4*cm, id="cl")],  onPage=_rne_on_closing),
    ])

    story: List[Any] = []
    ohx = _C_ORANGE.hexval()[2:]; mhx = _C_MUTED.hexval()[2:]
    ghx = _C_GREEN.hexval()[2:];  rhx = _C_RED.hexval()[2:]
    n_docs   = len(docs)
    n_bilans = company_summary.get("nb_comptes_annuels", company_summary.get("nb_bilans", 0))
    n_docs_juridiques = company_summary.get("nb_documents_juridiques", 0)
    n_actes  = company_summary.get("nb_actes", 0)

    # ── Couverture ──
    story.append(Spacer(1, 2 * cm))
    if LOGO_PATH.exists():
        story.append(RLImage(str(LOGO_PATH), width=5.5*cm, height=2.9*cm))
    story.append(Spacer(1, 2 * cm))
    story.append(RLPara("Rapport de collecte RNE", sty["RCvT"]))
    story.append(Spacer(1, 8 * mm))
    story.append(RLPara(_rx(denomination), sty["RCvS"]))
    story.append(Spacer(1, 3 * mm))
    story.append(RLPara(f"SIREN {siren}", sty["RCvM"]))
    story.append(Spacer(1, 1.8 * cm))
    for ml in [f"Date du rapport : {run_date}",
               f"Documents collect\u00e9s : {n_docs}",
               f"FICHIERS ANALYSES - Bilans : {n_bilans} | Documents juridique : {n_docs_juridiques} | Actes : {n_actes}",
               "Source : Registre National des Entreprises (INPI)"]:
        story.append(RLPara(ml, sty["RCvM"])); story.append(Spacer(1, 1.5 * mm))

    # ── Sommaire ──
    story.append(NextPageTemplate("content")); story.append(PageBreak())
    story.append(RLPara("Sommaire", sty["RTocT"]))
    story.append(HRFlowable(width="25%", thickness=3, color=_C_ORANGE, spaceBefore=0, spaceAfter=10*mm, hAlign="LEFT"))
    toc_lines: List[Any] = [(1, "A", "Fiche entreprise")]
    for sec_num, sec_title in rne.get("section_order", []):
        toc_lines.append((2, sec_num, sec_title))
    toc_lines.extend([
        (1, "B", "Documents d\u00e9pos\u00e9s"),
        (2, "B.1", "Liste des documents"),
        (2, "B.2", "Analyse synth\u00e9tique des documents"),
        (1, "C", "Informations de collecte"),
    ])
    for level, num, title in toc_lines:
        s = sty["RToc1"] if level == 1 else sty["RToc2"]
        clr = ohx if level == 1 else mhx
        story.append(RLPara(f"<font color='#{clr}'><b>{num}</b></font>&nbsp;&nbsp;&nbsp;{_rx(title)}", s))

    # ── Section A ──
    story.append(PageBreak())
    story.append(RLPara(f"<font color='#{ohx}'>A</font>&nbsp;&nbsp;Fiche entreprise", sty["RH1"]))
    story.append(HRFlowable(width="100%", thickness=1.5, color=_C_ORANGE, spaceBefore=0, spaceAfter=6*mm))
    attestation_url = f"https://data.inpi.fr/export/companies?format=pdf&ids=[%22{siren}%22]"
    story.append(
        RLPara(
            f"""<link href="{_rx(attestation_url)}"><u><font color="#0066CC" size="11">👉 URL Cliquable</font></u></link>""",
            sty["RBd"],
        )
    )
    story.append(Spacer(1, 2 * mm))
    sections = rne.get("sections") or []
    if not sections:
        story.append(RLPara("Aucune donn\u00e9e entreprise exploitable.", sty["RBd"]))
    for sec in sections:
        rows = sec.get("rows") or []
        if not rows:
            continue
        story.append(RLPara(_rx(sec.get("title") or "Section"), sty["RH2"]))
        story.append(_rne_kv_table(rows, sty, total_w=cw))
        story.append(Spacer(1, 2 * mm))

    # ── Section B ──
    story.append(PageBreak())
    story.append(RLPara(f"<font color='#{ohx}'>B</font>&nbsp;&nbsp;Documents d\u00e9pos\u00e9s au RNE", sty["RH1"]))
    story.append(HRFlowable(width="100%", thickness=1.5, color=_C_ORANGE, spaceBefore=0, spaceAfter=6*mm))

    # B — Synthèse des volumes
    story.append(RLPara("Synthèse des documents analysés", sty["RH2"]))
    sum_hdr = [RLPara("<b>Type de document</b>", sty["RTh"]), RLPara("<b>Nombre</b>", sty["RTh"])]
    sum_rows: List[Any] = [
        sum_hdr,
        [RLPara("Bilans déposés", sty["RTd"]), RLPara(str(company_summary.get("nb_bilans", n_bilans)), sty["RTd"])],
        [RLPara("Documents juridiques", sty["RTd"]), RLPara(str(company_summary.get("nb_documents_juridiques", n_docs_juridiques)), sty["RTd"])],
        [RLPara("Actes détectés", sty["RTd"]), RLPara(str(company_summary.get("nb_actes", n_actes)), sty["RTd"])],
    ]
    sum_tbl = Table(sum_rows, colWidths=[cw - 2.0 * cm, 2.0 * cm], hAlign="LEFT")
    sum_style: List[Any] = [
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("BACKGROUND", (0, 0), (-1, 0), _C_NAVY),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("BOX", (0, 0), (-1, -1), 0.3, _C_BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.3, _C_BORDER),
    ]
    for i in range(1, len(sum_rows)):
        sum_style.append(("BACKGROUND", (0, i), (-1, i), _C_LGRAY if i % 2 == 0 else _C_CARD))
    sum_tbl.setStyle(TableStyle(sum_style))
    story.append(sum_tbl)
    story.append(Spacer(1, 5 * mm))

    story.append(RLPara("B.1&nbsp;&nbsp;Liste des documents", sty["RH2"]))
    if not docs:
        story.append(RLPara("Aucun document trouv\u00e9.", sty["RBd"]))
    else:
        hdr_b = [RLPara(f"<b>{h}</b>", sty["RTh"]) for h in
                 ["#", "Date d\u00e9p\u00f4t", "Type", "Nature INPI", "Composition", "Analys\u00e9"]]
        b_rows: List[Any] = [hdr_b]
        for ix, r in enumerate(docs, 1):
            bdg = r["document_analyse"]
            a_s = sty["RBgO"] if bdg == "OUI" else sty["RBgN"]
            tl = r.get("typeLibelle") or r.get("typeDocument") or ""
            if str(r.get("famille") or "").upper().startswith("COMP") and (not tl or str(tl).strip().upper() in {"N/A", "NA"}):
                tl = "Bilan"
            nat = tl
            nb_sub = int(r.get("nb_actes_in_doc") or 1)
            composition = f"{nb_sub} acte(s)" if str(r.get("famille") or "").upper() == "ACTE" else "Document unique"
            b_rows.append([RLPara(str(ix), sty["RTd"]), RLPara(_rx(r["date_depot"]), sty["RTd"]),
                           RLPara(_rx(r["famille"]), sty["RTd"]), RLPara(_rx(nat), sty["RTd"]),
                           RLPara(_rx(composition), sty["RTd"]),
                           RLPara(bdg, a_s)])
        bw = [1.6*cm, 2.7*cm, 3.1*cm, 6.2*cm, cw-1.6*cm-2.7*cm-3.1*cm-6.2*cm-2*cm, 2*cm]
        bt = Table(b_rows, colWidths=bw)
        bts: List[Any] = [("VALIGN",(0,0),(-1,-1),"MIDDLE"),("TOPPADDING",(0,0),(-1,-1),5),
                ("BOTTOMPADDING",(0,0),(-1,-1),5),("LEFTPADDING",(0,0),(-1,-1),6),
                ("BACKGROUND",(0,0),(-1,0),_C_NAVY),("TEXTCOLOR",(0,0),(-1,0),white),
                ("LINEBELOW",(0,1),(-1,-1),0.3,_C_BORDER)]
        for i in range(1, len(b_rows)):
            bts.append(("BACKGROUND",(0,i),(-1,i),_C_LGRAY if i%2==0 else _C_CARD))
        bt.setStyle(TableStyle(bts)); story.append(bt)

    story.append(PageBreak())
    story.append(RLPara("B.2&nbsp;&nbsp;Analyse synth\u00e9tique des documents", sty["RH2"]))

    def _is_not_available(v: Any) -> bool:
        s = str(v or "").strip()
        return s == "" or s.upper() in {"N/A", "NA", "-", "\u2014"}

    def _display_type_label(row: Dict[str, Any]) -> str:
        tlib = str(row.get("typeLibelle") or "").strip()
        if not _is_not_available(tlib):
            return tlib
        labels = [str(x).strip() for x in (row.get("actes_labels") or []) if str(x).strip()]
        if labels:
            return " + ".join(labels[:4])
        tcode = str(row.get("typeDocument") or "").strip()
        if not _is_not_available(tcode):
            return tcode
        fam = str(row.get("famille") or "").upper()
        if fam.startswith("COMP"):
            return "Comptes annuels"
        if fam == "ACTE":
            return "Acte juridique"
        return "Document non typé"

    def _display_zip_path(row: Dict[str, Any]) -> str:
        fam = str(row.get("famille") or "").upper()
        prefix = "Actes/" if fam == "ACTE" else "Bilans/"
        return f"{prefix}{row.get('filename_base', 'Document')}.pdf"

    def _clean_descriptif(v: Any) -> str:
        s = str(v or "").strip()
        if not s:
            return ""
        chunks = [re.sub(r"\s+", " ", c).strip(" -|;") for c in re.split(r"\s*\|\s*|\s*;\s*", s) if c.strip()]
        out: List[str] = []
        seen = set()
        for c in chunks:
            k = c.casefold()
            if k in seen:
                continue
            seen.add(k)
            out.append(c)
        return "\n".join(out[:12])

    for ix, r in enumerate(docs, 1):
        story.append(Spacer(1, 4*mm))
        date_depot = str(r.get("date_depot") or "").strip() or "Date non disponible"
        display_type = _display_type_label(r)
        title_tbl = Table([[RLPara(f"<b>{_rx(date_depot)} / {_rx(display_type)}</b>", sty["RTh"])]], colWidths=[cw])
        title_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), _C_NAVY),
            ("TEXTCOLOR", (0, 0), (-1, -1), white),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(title_tbl)

        fam = str(r.get("famille") or "").strip() or "Document"
        status_raw = str(r.get("document_analyse") or "").upper()
        status_label = "OUI" if status_raw == "OUI" else "NON"
        cleaned_desc = _clean_descriptif(r.get("descriptif"))
        display_description = "Analyse non disponible pour ce document." if status_label == "NON" else (
            cleaned_desc if cleaned_desc else "Aucune description exploitable extraite."
        )
        meta_rows: List[Any] = [
            ("Nom du fichier", f"{r.get('filename_base', 'Document')}.pdf"),
            ("Famille", fam),
            ("Type de document", display_type),
            ("Statut d'analyse", status_label),
        ]
        if str(r.get("famille") or "").upper() == "ACTE":
            meta_rows.append(("Nombre d'actes détectés", str(int(r.get("nb_actes_in_doc") or 1))))
        meta_rows.append(("Description", display_description))
        story.append(_rne_kv_table(meta_rows, sty, col1=5.3*cm, total_w=cw))

        labels = [str(x).strip() for x in (r.get("actes_labels") or []) if str(x).strip()]
        if str(r.get("famille") or "").upper() == "ACTE" and labels:
            story.append(Spacer(1, 1.5 * mm))
            story.append(RLPara("Actes identifiés :", sty["RLbl"]))
            for lbl in labels:
                story.append(RLPara(f"- {_rx(lbl)}", sty["RBd"]))

        nb_sub = int(r.get("nb_actes_in_doc") or 1)
        if str(r.get("famille") or "").upper() == "ACTE":
            if nb_sub > 1:
                story.append(Spacer(1, 1.5 * mm))
                story.append(RLPara(f"Document composite comprenant {nb_sub} actes.", sty["RBd"]))

        if ix < n_docs:
            story.append(Spacer(1, 3*mm))
            story.append(HRFlowable(width="100%", thickness=0.3, color=_C_BORDER, spaceBefore=0, spaceAfter=2*mm))

    # ── Section C ──
    story.append(PageBreak())
    story.append(RLPara(f"<font color='#{ohx}'>C</font>&nbsp;&nbsp;Informations de collecte", sty["RH1"]))
    story.append(HRFlowable(width="100%", thickness=1.5, color=_C_ORANGE, spaceBefore=0, spaceAfter=6*mm))
    pts = rne.get("points_attention") or []
    if pts:
        apars = [RLPara(f"<font color='#{ohx}'>\u25b8</font>&nbsp;&nbsp;{_rx(pt)}", sty["RAlt"]) for pt in pts]
        story.append(_AlertBox(cw, apars, title="Informations de collecte"))
    else:
        story.append(RLPara("Aucun point d\u2019attention identifi\u00e9.", sty["RBd"]))

    # ── Page finale ──
    story.append(NextPageTemplate("closing")); story.append(PageBreak())
    story.append(Spacer(1, 2.5*cm))
    if LOGO_PATH.exists():
        lg = RLImage(str(LOGO_PATH), width=5.5*cm, height=2.9*cm); lg.hAlign = "CENTER"; story.append(lg)
    story.append(Spacer(1, 1.5*cm))
    story.append(RLPara("Automatiser la collecte RNE dans votre cabinet", sty["RClT"]))
    story.append(Spacer(1, 6*mm))
    story.append(RLPara(
        "Cet agent permet de collecter et structurer automatiquement les donn\u00e9es publiques "
        "du Registre National des Entreprises (INPI) afin de produire des rapports exploitables "
        "par les \u00e9quipes du cabinet.", sty["RClS"]))
    story.append(Spacer(1, 5*mm))
    story.append(RLPara("Nous \u00e9tudions avec vous :", sty["RClS"]))
    story.append(Spacer(1, 2*mm))
    for blt in ["l\u2019int\u00e9gration dans vos logiciels de production et de gestion interne",
                "l\u2019adaptation de l\u2019agent \u00e0 vos processus m\u00e9tier",
                "l\u2019automatisation de la collecte documentaire RNE"]:
        story.append(RLPara(f"\u2022  {blt}", sty["RClS"]))
    story.append(Spacer(1, 8*mm))
    story.append(RLPara("<b>Contact : contact@iagidoo.fr</b>", sty["RClS"]))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ── WORD helpers ──────────────────────────────────────────────────────────────
def _w_strip_md(t: str) -> str:
    import re as _re
    return _re.sub(r'\*\*(.*?)\*\*', r'\1', str(t))


def _w_dash(v: Any) -> str:
    s = str(v).strip() if v not in (None, "", [], {}) else ""
    return s if s else "\u2014"


def _w_set_cell_bg(cell: Any, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.replace("#", ""))
    tcPr.append(shd)


def _w_set_cell_border(cell: Any, **kwargs: Any) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        v = kwargs.get(side, {})
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), v.get("val", "none"))
        border.set(qn("w:sz"), str(v.get("sz", 0)))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), v.get("color", "auto"))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _w_add_heading1(doc: Any, letter: str, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run_l = p.add_run(letter + "  ")
    run_l.font.color.rgb = _D_ORANGE; run_l.font.size = Pt(20); run_l.font.bold = True
    run_t = p.add_run(title)
    run_t.font.color.rgb = _D_NAVY; run_t.font.size = Pt(20); run_t.font.bold = True
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(8)
    rhr = p2.add_run("_" * 90)
    rhr.font.color.rgb = _D_ORANGE; rhr.font.size = Pt(5)


def _w_add_heading2(doc: Any, num: str, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    run_n = p.add_run(num + "  ")
    run_n.font.color.rgb = _D_TEAL; run_n.font.size = Pt(13); run_n.font.bold = True
    run_t = p.add_run(title)
    run_t.font.color.rgb = _D_TEAL; run_t.font.size = Pt(13); run_t.font.bold = True


def _w_kv_table(doc: Any, rows: List[Any], col1_cm: float = 5.5, total_cm: float = 17.0) -> Any:
    col2_cm = total_cm - col1_cm
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = DocxCm(col1_cm)
    table.columns[1].width = DocxCm(col2_cm)
    for i, (lbl, val) in enumerate(rows):
        bg = "F3F1EE" if i % 2 == 0 else "FFFFFF"
        c0 = table.rows[i].cells[0]
        _w_set_cell_bg(c0, bg)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(3); p0.paragraph_format.space_after = Pt(3)
        r0 = p0.add_run(str(lbl)); r0.font.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = _D_MUTED
        c1 = table.rows[i].cells[1]
        _w_set_cell_bg(c1, bg)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(3); p1.paragraph_format.space_after = Pt(3)
        r1 = p1.add_run(str(val)); r1.font.size = Pt(9.5)
    for row in table.rows:
        for cell in row.cells:
            _w_set_cell_border(cell,
                bottom={"val": "single", "sz": 2, "color": "E8E4DF"},
                top={"val": "none"}, left={"val": "none"}, right={"val": "none"})
    return table


def _w_doc_table_header(doc: Any, headers: List[str], widths_cm: List[float]) -> Any:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (h, w) in enumerate(zip(headers, widths_cm)):
        table.columns[i].width = DocxCm(w)
        cell = table.rows[0].cells[i]
        _w_set_cell_bg(cell, "1A2744")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(h); r.font.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = _D_WHITE
    return table


def _w_doc_table_row(table: Any, values: List[Any], row_index: int) -> None:
    row = table.add_row()
    bg = "F3F1EE" if row_index % 2 == 0 else "FFFFFF"
    for i, val in enumerate(values):
        cell = row.cells[i]
        _w_set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        if isinstance(val, tuple):
            text, color = val
            r = p.add_run(str(text)); r.font.size = Pt(8.5); r.font.bold = True
            r.font.color.rgb = _D_GREEN if color == "green" else _D_RED
        else:
            r = p.add_run(str(val)); r.font.size = Pt(8.5)
        _w_set_cell_border(cell,
            bottom={"val": "single", "sz": 2, "color": "E8E4DF"},
            top={"val": "none"}, left={"val": "none"}, right={"val": "none"})


def _w_alert_box(doc: Any, points: List[str]) -> None:
    if not points:
        p = doc.add_paragraph("Aucun point d\u2019attention identifi\u00e9.")
        p.runs[0].font.size = Pt(9.5)
        return
    table = doc.add_table(rows=1 + len(points), cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = DocxCm(17.0)
    cell_hdr = table.rows[0].cells[0]
    _w_set_cell_bg(cell_hdr, "EF8829")
    p_hdr = cell_hdr.paragraphs[0]
    p_hdr.paragraph_format.space_before = Pt(4); p_hdr.paragraph_format.space_after = Pt(4)
    rh = p_hdr.add_run("Informations de collecte")
    rh.font.bold = True; rh.font.size = Pt(9.5); rh.font.color.rgb = _D_WHITE
    for i, pt in enumerate(points):
        cell = table.rows[i + 1].cells[0]
        _w_set_cell_bg(cell, "FFF8F0")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = DocxCm(0.3)
        rb = p.add_run("\u25b8  "); rb.font.color.rgb = _D_ORANGE; rb.font.size = Pt(9.5)
        rt = p.add_run(str(pt)); rt.font.size = Pt(9.5)
        _w_set_cell_border(cell,
            top={"val": "none"},
            left={"val": "single", "sz": 4, "color": "EF8829"},
            right={"val": "single", "sz": 4, "color": "EF8829"},
            bottom=({"val": "single", "sz": 2, "color": "E8E4DF"} if i < len(points) - 1
                    else {"val": "single", "sz": 4, "color": "EF8829"}))


# ── generate_report_word ───────────────────────────────────────────────────────
def generate_report_word(
    siren: str,
    denomination: str,
    company_summary: Dict[str, Any],
    doc_results: List[Dict[str, Any]],
    run_date: str,
) -> bytes:
    rne_data = _to_rne_data(company_summary, siren, denomination)
    norm_docs = _norm_doc_results(doc_results)

    wdoc = DocxDocument()
    for section in wdoc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = DocxCm(29.7)
        section.page_height = DocxCm(21.0)
        section.top_margin = DocxCm(2.0); section.bottom_margin = DocxCm(2.0)
        section.left_margin = DocxCm(2.5); section.right_margin = DocxCm(2.5)

    n_docs = len(norm_docs)
    n_bilans = sum(1 for d in norm_docs if (d.get("famille") or "").upper().startswith("COMP"))
    n_docs_juridiques = company_summary.get("nb_documents_juridiques", 0)
    n_actes = sum(1 for d in norm_docs if (d.get("famille") or "").upper() == "ACTE")

    # ── Couverture ────────────────────────────────────────────────────────────
    if LOGO_PATH.exists():
        p_logo = wdoc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(str(LOGO_PATH), width=DocxCm(6))
    else:
        pl = wdoc.add_paragraph("GIDOO")
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl.runs[0].font.size = Pt(28); pl.runs[0].font.bold = True

    wdoc.add_paragraph()

    p_title = wdoc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rt = p_title.add_run("Rapport de collecte RNE")
    rt.font.size = Pt(32); rt.font.bold = True; rt.font.color.rgb = _D_NAVY

    p_sub = wdoc.add_paragraph()
    rs = p_sub.add_run(denomination)
    rs.font.size = Pt(16); rs.font.color.rgb = _D_TEAL

    p_siren = wdoc.add_paragraph()
    rsr = p_siren.add_run(f"SIREN {siren}")
    rsr.font.size = Pt(11); rsr.font.color.rgb = _D_MUTED

    p_sep = wdoc.add_paragraph("_" * 90)
    p_sep.runs[0].font.color.rgb = _D_ORANGE; p_sep.runs[0].font.size = Pt(5)

    for ml in [
        f"Date du rapport : {run_date}",
        f"Documents collect\u00e9s : {n_docs}",
        f"FICHIERS ANALYSES - Bilans : {n_bilans} | Documents juridique : {n_docs_juridiques} | Actes : {n_actes}",
        "Source : Registre National des Entreprises (INPI)",
    ]:
        pm = wdoc.add_paragraph()
        rm = pm.add_run(ml); rm.font.size = Pt(10); rm.font.color.rgb = _D_MUTED

    # ── Sommaire ──────────────────────────────────────────────────────────────
    wdoc.add_page_break()
    p_toc = wdoc.add_paragraph()
    rtoc = p_toc.add_run("Sommaire")
    rtoc.font.size = Pt(22); rtoc.font.bold = True; rtoc.font.color.rgb = _D_NAVY
    p_toc.paragraph_format.space_after = Pt(8)

    p_sep2 = wdoc.add_paragraph("_" * 30)
    p_sep2.runs[0].font.color.rgb = _D_ORANGE; p_sep2.runs[0].font.size = Pt(5)
    p_sep2.paragraph_format.space_after = Pt(12)

    toc_lines: List[Any] = [(1, "A", "Fiche entreprise")]
    for sec_num, sec_title in rne_data.get("section_order", []):
        toc_lines.append((2, sec_num, sec_title))
    toc_lines.extend([
        (1, "B", "Documents d\u00e9pos\u00e9s"),
        (2, "B.1", "Liste des documents"),
        (2, "B.2", "D\u00e9tail des documents"),
        (1, "C", "Informations de collecte"),
    ])
    for level, num, ttl in toc_lines:
        pi = wdoc.add_paragraph()
        pi.paragraph_format.space_before = Pt(0); pi.paragraph_format.space_after = Pt(0)
        if level == 2:
            pi.paragraph_format.left_indent = DocxCm(0.8)
        rn = pi.add_run(f"{num}   ")
        rn.font.bold = True; rn.font.size = Pt(11) if level == 1 else Pt(10)
        rn.font.color.rgb = _D_ORANGE if level == 1 else _D_MUTED
        rt2 = pi.add_run(ttl)
        rt2.font.size = Pt(11) if level == 1 else Pt(10); rt2.font.bold = (level == 1)

    # ── Section A ─────────────────────────────────────────────────────────────
    wdoc.add_page_break()
    _w_add_heading1(wdoc, "A", "Fiche entreprise")
    sections = rne_data.get("sections") or []
    if not sections:
        p = wdoc.add_paragraph("Aucune donn\u00e9e entreprise exploitable.")
        p.runs[0].font.size = Pt(9.5)
    for sec in sections:
        rows = sec.get("rows") or []
        if not rows:
            continue
        title = str(sec.get("title") or "Section")
        if " " in title:
            num, ttl = title.split(" ", 1)
        else:
            num, ttl = title, title
        _w_add_heading2(wdoc, num.strip(), ttl.strip())
        _w_kv_table(wdoc, rows)

    # ── Section B ─────────────────────────────────────────────────────────────
    wdoc.add_page_break()
    _w_add_heading1(wdoc, "B", "Documents d\u00e9pos\u00e9s")

    _w_add_heading2(wdoc, "B.1", "Liste des documents")
    if not norm_docs:
        p = wdoc.add_paragraph("Aucun document trouv\u00e9."); p.runs[0].font.size = Pt(9.5)
    else:
        tbl_docs = _w_doc_table_header(wdoc,
            ["#", "Date d\u00e9p\u00f4t", "Type", "Nature INPI", "Composition", "Analys\u00e9"],
            [0.9, 2.4, 2.8, 5.8, 4.1, 2.0])
        for ix, r in enumerate(norm_docs, 1):
            analysed = r.get("document_analyse", "NON")
            tl = r.get("typeLibelle") or r.get("nature") or r.get("typeDocument") or ""
            nature = tl
            nb_sub = int(r.get("nb_actes_in_doc") or 1)
            composition = f"{nb_sub} acte(s)" if str(r.get("famille") or "").upper() == "ACTE" else "Document unique"
            _w_doc_table_row(tbl_docs,
                [str(ix), r.get("date_depot", ""), r.get("famille", ""),
                 nature, composition, (analysed, "green" if analysed == "OUI" else "red")],
                ix - 1)

    _w_add_heading2(wdoc, "B.2", "D\u00e9tail des documents")
    for ix, r in enumerate(norm_docs, 1):
        wdoc.add_paragraph()
        bdg = r.get("document_analyse", "NON")
        fname = r.get("filename_base", r.get("filename", f"Document_{ix}"))
        p_dt = wdoc.add_paragraph()
        p_dt.paragraph_format.space_before = Pt(6)
        rnum = p_dt.add_run(f"Document {ix}/{n_docs}  ")
        rnum.font.bold = True; rnum.font.size = Pt(11); rnum.font.color.rgb = _D_NAVY
        rbdg = p_dt.add_run(f"[{bdg}]  ")
        rbdg.font.bold = True; rbdg.font.size = Pt(11)
        rbdg.font.color.rgb = _D_GREEN if bdg == "OUI" else _D_RED
        rnm = p_dt.add_run(f"{fname}.pdf")
        rnm.font.size = Pt(11); rnm.font.color.rgb = _D_NAVY

        tl = r.get("typeLibelle") or r.get("nature") or r.get("typeDocument") or ""
        detail_rows = [
            ("Date de d\u00e9p\u00f4t", r.get("date_depot", "")),
            ("Famille",                 r.get("famille", "")),
            ("Nature INPI",             tl),
        ]
        nb_sub = int(r.get("nb_actes_in_doc") or 1)
        if str(r.get("famille") or "").upper() == "ACTE":
            if nb_sub > 1:
                detail_rows.append(("Document composite", f"Document composite comportant {nb_sub} actes"))
            else:
                detail_rows.append(("Document composite", "Document comportant 1 acte"))
        _w_kv_table(wdoc, detail_rows, col1_cm=4.0)
        labels = [str(x).strip() for x in (r.get("actes_labels") or []) if str(x).strip()]
        if labels:
            p_lbl_act = wdoc.add_paragraph()
            p_lbl_act.paragraph_format.space_before = Pt(4)
            rl_act = p_lbl_act.add_run("Actes détectés")
            rl_act.font.bold = True
            rl_act.font.size = Pt(9)
            for act_lbl in labels:
                p_a = wdoc.add_paragraph(f"•  {act_lbl}")
                p_a.runs[0].font.size = Pt(9.5)

        descriptif = r.get("descriptif") or r.get("description") or ""
        if descriptif:
            p_lbl = wdoc.add_paragraph()
            p_lbl.paragraph_format.space_before = Pt(4)
            rl = p_lbl.add_run("Descriptif")
            rl.font.bold = True; rl.font.size = Pt(9); rl.font.color.rgb = _D_MUTED
            for dl in _w_strip_md(descriptif.strip()).split("\n"):
                ds = dl.strip()
                if ds:
                    pd = wdoc.add_paragraph(ds); pd.runs[0].font.size = Pt(9.5)
            texte = r.get("texte_extrait") or r.get("texte") or ""
            if texte:
                p_te_lbl = wdoc.add_paragraph()
                rte = p_te_lbl.add_run("TEXTE EXTRAIT :")
                rte.font.bold = True; rte.font.size = Pt(9); rte.font.color.rgb = _D_MUTED
                excerpt = texte.strip()[:600] + ("\u2026" if len(texte.strip()) > 600 else "")
                pte = wdoc.add_paragraph(excerpt)
                pte.runs[0].font.size = Pt(8.5); pte.runs[0].font.color.rgb = _D_MUTED
        else:
            pnd = wdoc.add_paragraph("Aucun descriptif disponible.")
            pnd.runs[0].font.italic = True; pnd.runs[0].font.size = Pt(8.5)
            pnd.runs[0].font.color.rgb = _D_MUTED
        if ix < n_docs:
            psd = wdoc.add_paragraph("_" * 90)
            psd.runs[0].font.color.rgb = RGBColor(0xE8, 0xE4, 0xDF); psd.runs[0].font.size = Pt(4)

    # ── Section C ─────────────────────────────────────────────────────────────
    wdoc.add_page_break()
    _w_add_heading1(wdoc, "C", "Informations de collecte")
    _w_alert_box(wdoc, rne_data.get("points_attention", []))

    # ── Page finale ───────────────────────────────────────────────────────────
    wdoc.add_page_break()
    if LOGO_PATH.exists():
        p_logo2 = wdoc.add_paragraph()
        p_logo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo2.add_run().add_picture(str(LOGO_PATH), width=DocxCm(6))

    p_cl = wdoc.add_paragraph()
    p_cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cl.paragraph_format.space_before = Pt(20)
    rcl = p_cl.add_run("Automatiser la collecte RNE dans votre cabinet")
    rcl.font.size = Pt(16); rcl.font.bold = True; rcl.font.color.rgb = _D_NAVY

    p_cl2 = wdoc.add_paragraph()
    p_cl2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rcl2 = p_cl2.add_run(
        "Cet agent permet de collecter et structurer automatiquement les donn\u00e9es publiques "
        "du Registre National des Entreprises (INPI) afin de produire des rapports exploitables "
        "par les \u00e9quipes du cabinet.")
    rcl2.font.size = Pt(10); rcl2.font.color.rgb = _D_MUTED

    p_nous = wdoc.add_paragraph()
    p_nous.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nous.paragraph_format.space_before = Pt(10)
    rnous = p_nous.add_run("Nous \u00e9tudions avec vous :")
    rnous.font.size = Pt(10); rnous.font.color.rgb = _D_MUTED

    for blt in [
        "l\u2019int\u00e9gration dans vos logiciels de production et de gestion interne",
        "l\u2019adaptation de l\u2019agent \u00e0 vos processus m\u00e9tier",
        "l\u2019automatisation de la collecte documentaire RNE",
    ]:
        pb = wdoc.add_paragraph(f"\u2022  {blt}")
        pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pb.runs[0].font.size = Pt(10); pb.runs[0].font.color.rgb = _D_MUTED

    p_contact = wdoc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(14)
    rcon = p_contact.add_run("Contact : contact@iagidoo.fr")
    rcon.font.size = Pt(11); rcon.font.bold = True; rcon.font.color.rgb = _D_TEAL

    buf_w = io.BytesIO()
    wdoc.save(buf_w)
    buf_w.seek(0)
    return buf_w.getvalue()


def run_collection(job: Job) -> None:
    try:
        ok_cfg, missing = validate_config()
        if not ok_cfg:
            msg = f"Configuration INPI manquante: {', '.join(missing)}"
            job.errors.append(msg)
            job.status = "error"
            job.log(msg, 100)
            return

        job.status = "running"
        job.log("Authentification INPI...", 5)
        token = inpi_login()
        job.log("Authentification OK", 10)

        job.log(f"Lecture données entreprise SIREN {job.siren}...", 12)
        company_json = {}
        try:
            company_json = inpi_get_company(job.siren, token)
        except Exception as e:
            job.log(f"Avertissement: données entreprise non disponibles ({e})")

        job.log("Récupération des pièces...", 15)
        data = inpi_get_attachments(job.siren, token)
        bilans = data.get("bilans", [])
        actes = data.get("actes", [])

        if not bilans and not actes:
            job.errors.append("Aucun document trouvé")
            job.status = "completed"
            job.log("Aucun document trouvé", 100)
            return

        job.total_docs = len(bilans) + len(actes)
        job.estimated_minutes = max(1, round(job.total_docs * 0.35))

        company_summary = build_company_summary(company_json, job.siren, bilans, actes)
        job.denomination = (company_summary.get("denomination", "") or "").strip() or (
            (bilans[0].get("denomination", "") if bilans else (actes[0].get("denomination", "") if actes else ""))
        )

        job.log("Analyse entreprise structurée générée", 22)

        zip_buf = io.BytesIO()
        doc_results: List[Dict[str, Any]] = []
        run_date = datetime.now().strftime("%Y-%m-%d")

        ppd = 60 / max(job.total_docs, 1)
        cp = 25

        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            used_zip_paths: set[str] = set()
            for i, b in enumerate(bilans):
                bid = b.get("id")
                if not bid:
                    continue
                cp += ppd
                job.log(f"Bilan {i + 1}/{len(bilans)}...", int(cp))
                try:
                    content = inpi_download(f"bilans/{bid}/download", token)
                    res = process_one_document("COMPTES ANNUELS", b, content, job.siren, job.denomination)
                    res["endpoint"] = f"bilans/{bid}/download"
                    bil_path = unique_zip_path(f"COMPTES_ANNUELS/{res['filename_base']}.pdf", used_zip_paths)
                    zf.writestr(bil_path, content)
                    rows = res.get("report_rows", [res])
                    for r in rows:
                        r["endpoint"] = res["endpoint"]
                    doc_results.extend(rows)
                    job.bilans_count += 1
                    job.current_doc_name = res["filename_base"]
                    job.current_doc_desc = res["descriptif"][:300]
                except Exception as e:
                    job.errors.append(f"Bilan {bid}: {e}")
                time.sleep(0.1)

            for i, a in enumerate(actes):
                aid = a.get("id")
                if not aid:
                    continue
                cp += ppd
                job.log(f"Acte {i + 1}/{len(actes)}...", int(cp))
                try:
                    ai = inpi_get_acte_info(aid, token)
                    merged = {**a, **ai}
                    content = inpi_download(f"actes/{aid}/download", token)
                    res = process_one_document("ACTE", merged, content, job.siren, job.denomination)
                    res["endpoint"] = f"actes/{aid}/download"
                    act_path = unique_zip_path(f"ACTES/{res['filename_base']}.pdf", used_zip_paths)
                    zf.writestr(act_path, content)
                    rows = res.get("report_rows", [res])
                    for r in rows:
                        r["endpoint"] = res["endpoint"]
                    doc_results.extend(rows)
                    job.actes_count += 1
                    job.current_doc_name = res["filename_base"]
                    job.current_doc_desc = res["descriptif"][:300]
                except Exception as e:
                    job.errors.append(f"Acte {aid}: {e}")
                time.sleep(0.1)

            try:
                company_pdf = inpi_download_company_export_pdf(job.siren)
                if company_pdf:
                    export_name = make_attestation_filename(run_date, job.siren, job.denomination or "Entreprise")
                    export_path = unique_zip_path(f"{export_name}.pdf", used_zip_paths)
                    zf.writestr(export_path, company_pdf)
            except Exception:
                pass

            # Recompute counters including composite documents (typeRdd => n actes métier).
            total_actes = sum(
                int(d.get("nb_actes_in_doc") or 1)
                for d in doc_results
                if str(d.get("source", "")).upper() == "ACTE"
            )
            company_summary["nb_comptes_annuels"] = len(bilans)
            company_summary["nb_documents_juridiques"] = len(actes)
            company_summary["nb_actes"] = total_actes
            company_summary["nb_docs"] = len(bilans) + len(actes)
            job.actes_count = total_actes

            job.log("Génération rapport PDF...", 90)
            report_bytes = generate_report_pdf(
                job.siren,
                job.denomination or "Entreprise",
                company_summary,
                doc_results,
                run_date,
            )
            report_name = make_report_filename(run_date, job.siren, job.denomination or "Entreprise")
            report_pdf_path = unique_zip_path(f"{report_name}.pdf", used_zip_paths)
            zf.writestr(report_pdf_path, report_bytes)
            word_bytes = generate_report_word(
                job.siren,
                job.denomination or "Entreprise",
                company_summary,
                doc_results,
                run_date,
            )
            report_docx_path = unique_zip_path(f"{report_name}.docx", used_zip_paths)
            zf.writestr(report_docx_path, word_bytes)

        zip_buf.seek(0)
        job.zip_data = zip_buf.read()
        job.log(f"Terminé: {job.bilans_count} bilan(s), {job.actes_count} acte(s)", 100)
        job.status = "completed"

    except requests.HTTPError as e:
        msg = f"Erreur API INPI: {e.response.status_code if e.response is not None else 'HTTP'}"
        job.errors.append(msg)
        job.log(msg, 100)
        job.status = "error"
    except Exception as e:
        job.errors.append(str(e))
        job.log(str(e), 100)
        job.status = "error"


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/collect", methods=["POST"])
def api_collect():
    data = request.json or {}
    siren = re.sub(r"\s+", "", data.get("siren", ""))
    if not re.match(r"^\d{9}$", siren):
        return jsonify({"error": "SIREN invalide: 9 chiffres requis"}), 400

    job = Job(siren=siren)
    with jobs_lock:
        jobs[job.id] = job

    threading.Thread(target=run_collection, args=(job,), daemon=True).start()
    return jsonify({"job_id": job.id})


@app.route("/api/status/<job_id>")
def api_status(job_id: str):
    with jobs_lock:
        job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job non trouvé"}), 404
    return jsonify(job.to_dict())


@app.route("/api/download/<job_id>")
def api_download(job_id: str):
    with jobs_lock:
        job = jobs.get(job_id)
    if not job or not job.zip_data:
        return jsonify({"error": "Aucun ZIP disponible"}), 404

    fn = sanitize_filename(f"{job.siren} - {datetime.now().strftime('%Y-%m-%d')}.zip")
    return send_file(io.BytesIO(job.zip_data), mimetype="application/zip", as_attachment=True, download_name=fn)


@app.route("/api/config")
def api_config():
    ok, missing = validate_config()
    return jsonify({"ok": ok, "missing": missing, "has_env_file": (BASE_DIR / ".env").exists()})


@app.route("/api/version")
def api_version():
    return jsonify(
        {
            "version": f"vscode-rne-{datetime.now().strftime('%Y%m%d')}",
            "python": sys.version.split(" ")[0],
            "fitz": fitz is not None,
            "openpyxl": load_workbook is not None,
        }
    )


@app.route("/health")
def health():
    return jsonify({"ok": True, "service": "agent-rne-vscode", "time": datetime.now().isoformat()})


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=PORT, debug=False)

