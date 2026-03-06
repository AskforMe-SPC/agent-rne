"""
generate_report.py — Générateur de rapports PDF et Word (DOCX) pour l'analyse RNE Gidoo.

Usage:
    python generate_report.py --siren 983670290

Génère:
    - <date>_<siren>_<denomination>_Rapport_Analyse_RNE.pdf
    - <date>_<siren>_<denomination>_Rapport_Analyse_RNE.docx
"""

import io
import os
import re
import json
import base64 as _b64
from datetime import datetime

# ─── ReportLab (PDF) ──────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm, cm
from reportlab.lib.colors import HexColor, white, black
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    BaseDocTemplate, Frame, PageTemplate, NextPageTemplate,
    Paragraph as RLPara, Spacer, Table, TableStyle, PageBreak,
    Image as RLImage, HRFlowable, Flowable, KeepTogether,
)

# ─── python-docx (Word) ───────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════════════════════════
# PALETTE DE COULEURS
# ══════════════════════════════════════════════════════════════════════════════
C_ORANGE   = HexColor("#EF8829")
C_TEAL     = HexColor("#017e84")
C_NAVY     = HexColor("#1a2744")
C_NAVY_L   = HexColor("#243456")
C_TXT      = HexColor("#2c2c2c")
C_MUTED    = HexColor("#7a7470")
C_BORDER   = HexColor("#e8e4df")
C_LGRAY    = HexColor("#f3f1ee")
C_CARD     = HexColor("#ffffff")
C_RED      = HexColor("#c0392b")
C_GREEN    = HexColor("#27ae60")
C_ALERT_BG = HexColor("#fff8f0")
C_BG       = HexColor("#FAF8F6")

# DOCX colours (RGBColor)
D_ORANGE = RGBColor(0xEF, 0x88, 0x29)
D_TEAL   = RGBColor(0x01, 0x7E, 0x84)
D_NAVY   = RGBColor(0x1A, 0x27, 0x44)
D_MUTED  = RGBColor(0x7A, 0x74, 0x70)
D_GREEN  = RGBColor(0x27, 0xAE, 0x60)
D_RED    = RGBColor(0xC0, 0x39, 0x2B)
D_LGRAY  = RGBColor(0xF3, 0xF1, 0xEE)
D_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE SETUP (PDF)
# ══════════════════════════════════════════════════════════════════════════════
PAGE = A4
PW, PH = A4

LOGO_PATH = os.path.join(os.path.dirname(__file__), "assets", "logo.png")
if not os.path.exists(LOGO_PATH):
    LOGO_PATH = os.path.join(os.path.dirname(__file__), "static", "logo.png")

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def _x(t):
    """Escape XML pour ReportLab."""
    if not t:
        return ""
    return str(t).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _strip_md(t):
    """Supprime les marqueurs markdown bold."""
    return re.sub(r'\*\*', '', t) if t else t


def _logo_rl(w_cm=5.5):
    """Charge le logo pour ReportLab."""
    if os.path.exists(LOGO_PATH):
        return RLImage(LOGO_PATH, width=w_cm * cm, height=w_cm * cm * 0.50)
    return None


def _dash(v):
    """Retourne '—' si la valeur est vide/None."""
    return v if v and str(v).strip() and str(v).strip() != "—" else "—"


# ══════════════════════════════════════════════════════════════════════════════
# PDF — STYLES
# ══════════════════════════════════════════════════════════════════════════════
def _pdf_styles():
    s = getSampleStyleSheet()
    a = s.add

    # Cover
    a(ParagraphStyle("CvT", fontName="Helvetica-Bold", fontSize=34, leading=42,
                     textColor=white, alignment=TA_LEFT))
    a(ParagraphStyle("CvS", fontName="Helvetica", fontSize=15, leading=22,
                     textColor=HexColor("#b0c4de"), alignment=TA_LEFT))
    a(ParagraphStyle("CvM", fontName="Helvetica", fontSize=11, leading=16,
                     textColor=HexColor("#8899bb"), alignment=TA_LEFT))

    # TOC
    a(ParagraphStyle("TocT", fontName="Helvetica-Bold", fontSize=22, leading=28,
                     textColor=C_NAVY, spaceAfter=16))
    a(ParagraphStyle("Toc1", fontName="Helvetica-Bold", fontSize=11, leading=20,
                     textColor=C_TXT))
    a(ParagraphStyle("Toc2", fontName="Helvetica", fontSize=10, leading=18,
                     textColor=C_MUTED, leftIndent=18))

    # Headings
    a(ParagraphStyle("H1", fontName="Helvetica-Bold", fontSize=18, leading=24,
                     textColor=C_NAVY, spaceAfter=8))
    a(ParagraphStyle("H2", fontName="Helvetica-Bold", fontSize=13, leading=18,
                     textColor=C_TEAL, spaceBefore=12, spaceAfter=6))

    # Body
    a(ParagraphStyle("Bd", fontName="Helvetica", fontSize=9.5, leading=13.5,
                     textColor=C_TXT, alignment=TA_JUSTIFY, spaceAfter=4))
    a(ParagraphStyle("BdS", fontName="Helvetica", fontSize=8.5, leading=12,
                     textColor=C_TXT, spaceAfter=2))
    a(ParagraphStyle("Lbl", fontName="Helvetica-Bold", fontSize=9, leading=12,
                     textColor=C_MUTED))
    a(ParagraphStyle("Val", fontName="Helvetica", fontSize=9.5, leading=13,
                     textColor=C_TXT))

    # Table
    a(ParagraphStyle("Th", fontName="Helvetica-Bold", fontSize=8.5, leading=11,
                     textColor=white))
    a(ParagraphStyle("Td", fontName="Helvetica", fontSize=8.5, leading=12,
                     textColor=C_TXT))
    a(ParagraphStyle("TdB", fontName="Helvetica-Bold", fontSize=8.5, leading=12,
                     textColor=C_TXT))

    # Badges
    a(ParagraphStyle("BgO", fontName="Helvetica-Bold", fontSize=8.5, leading=11,
                     textColor=C_GREEN))
    a(ParagraphStyle("BgN", fontName="Helvetica-Bold", fontSize=8.5, leading=11,
                     textColor=C_RED))

    # Alert
    a(ParagraphStyle("Alt", fontName="Helvetica", fontSize=9.5, leading=14,
                     textColor=C_TXT, leftIndent=14))

    # Doc detail
    a(ParagraphStyle("DocT", fontName="Helvetica-Bold", fontSize=11, leading=15,
                     textColor=C_NAVY, spaceBefore=8, spaceAfter=4))

    # Closing page
    a(ParagraphStyle("ClT", fontName="Helvetica-Bold", fontSize=16, leading=22,
                     textColor=white, alignment=TA_CENTER))
    a(ParagraphStyle("ClS", fontName="Helvetica", fontSize=10, leading=14,
                     textColor=HexColor("#8899bb"), alignment=TA_CENTER))

    return s


# ══════════════════════════════════════════════════════════════════════════════
# PDF — PAGE CALLBACKS
# ══════════════════════════════════════════════════════════════════════════════
def _on_cover(c, doc):
    c.saveState()
    # Fond bleu marine
    c.setFillColor(C_NAVY)
    c.rect(0, 0, PW, PH, fill=1, stroke=0)
    # Barre orange en haut
    c.setFillColor(C_ORANGE)
    c.rect(0, PH - 7 * mm, PW, 7 * mm, fill=1, stroke=0)
    # Bloc décoratif en bas à droite
    c.setFillColor(C_NAVY_L)
    c.rect(PW - 90 * mm, 0, 90 * mm, 35 * mm, fill=1, stroke=0)
    # Ligne orange séparatrice
    c.setStrokeColor(C_ORANGE)
    c.setLineWidth(1.5)
    c.line(2.5 * cm, 7 * cm, PW - 2.5 * cm, 7 * cm)
    c.restoreState()


def _on_closing(c, doc):
    c.saveState()
    c.setFillColor(C_TEAL)
    c.rect(0, 0, PW, PH, fill=1, stroke=0)
    c.setFillColor(C_ORANGE)
    c.rect(0, 0, PW, 5 * mm, fill=1, stroke=0)
    c.restoreState()


def _on_content(c, doc):
    c.saveState()
    # Fond crème
    c.setFillColor(C_BG)
    c.rect(0, 0, PW, PH, fill=1, stroke=0)
    # En-tête
    c.setStrokeColor(C_ORANGE)
    c.setLineWidth(1)
    c.line(1.5 * cm, PH - 1.0 * cm, PW - 1.5 * cm, PH - 1.0 * cm)
    c.setFont("Helvetica", 6.5)
    c.setFillColor(C_MUTED)
    c.drawString(1.5 * cm, PH - 0.85 * cm, "Gidoo \u2014 Rapport d\u2019analyse RNE")
    dn = getattr(doc, '_dn', '')
    sr = getattr(doc, '_sr', '')
    c.drawRightString(PW - 1.5 * cm, PH - 0.85 * cm, f"{dn} \u2014 SIREN {sr}")
    # Pied de page
    c.setStrokeColor(C_BORDER)
    c.setLineWidth(0.4)
    c.line(1.5 * cm, 1.2 * cm, PW - 1.5 * cm, 1.2 * cm)
    c.setFont("Helvetica", 6.5)
    c.setFillColor(C_MUTED)
    rd = getattr(doc, '_rd', '')
    c.drawString(1.5 * cm, 0.8 * cm,
                 f"Rapport du {rd} \u2014 Source : INPI / Registre National des Entreprises")
    c.drawRightString(PW - 1.5 * cm, 0.8 * cm, f"Page {c.getPageNumber()}")
    c.restoreState()


# ══════════════════════════════════════════════════════════════════════════════
# PDF — ALERT BOX (Flowable)
# ══════════════════════════════════════════════════════════════════════════════
class AlertBox(Flowable):
    def __init__(self, width, paras, title="Points d\u2019attention"):
        Flowable.__init__(self)
        self.bw = width
        self.paras = paras
        self.title = title
        self._pad = 12
        aw = width - 2 * self._pad
        self._hs = []
        for p in paras:
            _, h = p.wrap(aw, 10000)
            self._hs.append(h)
        self._th = 26
        self.height = self._th + sum(self._hs) + 2 * self._pad + len(paras) * 3
        self.width = width

    def draw(self):
        c = self.canv
        p, r = self._pad, 5
        w, h = self.bw, self.height
        # Fond
        c.setFillColor(C_ALERT_BG)
        c.setStrokeColor(C_ORANGE)
        c.setLineWidth(0.5)
        c.roundRect(0, 0, w, h, r, fill=1, stroke=1)
        # Barre de titre orange arrondie en haut
        pth = c.beginPath()
        pth.moveTo(0, h - self._th)
        pth.lineTo(0, h - r)
        pth.arcTo(0, h - 2 * r, 2 * r, h, startAng=90, extent=90)
        pth.lineTo(w - r, h)
        pth.arcTo(w - 2 * r, h - 2 * r, w, h, startAng=0, extent=90)
        pth.lineTo(w, h - self._th)
        pth.close()
        c.setFillColor(C_ORANGE)
        c.drawPath(pth, fill=1, stroke=0)
        c.setFillColor(white)
        c.setFont("Helvetica-Bold", 9.5)
        c.drawString(p, h - self._th + 7, self.title)
        y = h - self._th - p
        aw = w - 2 * p
        for para in self.paras:
            _, ph = para.wrap(aw, 10000)
            para.drawOn(c, p, y - ph)
            y -= ph + 3


# ══════════════════════════════════════════════════════════════════════════════
# PDF — TABLE CLÉ/VALEUR
# ══════════════════════════════════════════════════════════════════════════════
def _kv_table(rows, sty, col1=5.5 * cm, total_w=None):
    if total_w is None:
        total_w = PW - 3.5 * cm
    col2 = total_w - col1
    data = []
    for lbl, val in rows:
        data.append([
            RLPara(f"<b>{_x(lbl)}</b>", sty["Lbl"]),
            RLPara(_x(str(val)), sty["Val"]),
        ])
    t = Table(data, colWidths=[col1, col2])
    ts = [
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING",   (0, 0), (-1, -1), 10),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 10),
        ("LINEBELOW",     (0, 0), (-1, -2), 0.3, C_BORDER),
    ]
    for i in range(len(data)):
        bg = C_LGRAY if i % 2 == 0 else C_CARD
        ts.append(("BACKGROUND", (0, i), (-1, i), bg))
    t.setStyle(TableStyle(ts))
    return t


# ══════════════════════════════════════════════════════════════════════════════
# GÉNÉRATION PDF
# ══════════════════════════════════════════════════════════════════════════════
def generate_report_pdf(siren, denomination, rne_data, doc_results, run_date):
    """
    Génère le rapport d'analyse RNE au format PDF.

    Args:
        siren (str): Numéro SIREN
        denomination (str): Nom de l'entreprise
        rne_data (dict): Données structurées de l'entreprise (voir format ci-dessous)
        doc_results (list): Liste des documents analysés
        run_date (str): Date du rapport (YYYY-MM-DD)

    Returns:
        bytes: Contenu du fichier PDF
    """
    buf = io.BytesIO()
    sty = _pdf_styles()

    doc = BaseDocTemplate(
        buf, pagesize=PAGE,
        topMargin=1.5 * cm, bottomMargin=1.5 * cm,
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
    )
    doc._dn = denomination
    doc._sr = siren
    doc._rd = run_date

    cw = PW - 3.0 * cm  # largeur de contenu

    doc.addPageTemplates([
        PageTemplate('cover',
                     frames=[Frame(2.5 * cm, 2 * cm, PW - 5 * cm, PH - 4 * cm, id='cf')],
                     onPage=_on_cover),
        PageTemplate('content',
                     frames=[Frame(1.5 * cm, 1.5 * cm, cw, PH - 3.2 * cm, id='ct')],
                     onPage=_on_content),
        PageTemplate('closing',
                     frames=[Frame(2 * cm, 2 * cm, PW - 4 * cm, PH - 4 * cm, id='cl')],
                     onPage=_on_closing),
    ])

    story = []
    _ohx = C_ORANGE.hexval()[2:]
    _mhx = C_MUTED.hexval()[2:]
    _ghx = C_GREEN.hexval()[2:]
    _rhx = C_RED.hexval()[2:]

    n_docs   = len(doc_results)
    n_bilans = sum(1 for d in doc_results if (d.get("famille") or "").upper().startswith("COMP"))
    n_actes  = sum(1 for d in doc_results if (d.get("famille") or "").upper() == "ACTE")

    # ─── PAGE 1 : COUVERTURE ───────────────────────────────────────────────
    story.append(Spacer(1, 2 * cm))
    logo = _logo_rl(5.5)
    if logo:
        story.append(logo)
    story.append(Spacer(1, 2 * cm))
    story.append(RLPara("Rapport d\u2019analyse RNE", sty["CvT"]))
    story.append(Spacer(1, 8 * mm))
    story.append(RLPara(_x(denomination), sty["CvS"]))
    story.append(Spacer(1, 3 * mm))
    story.append(RLPara(f"SIREN {siren}", sty["CvM"]))
    story.append(Spacer(1, 1.8 * cm))
    for ml in [
        f"Date du rapport : {run_date}",
        f"Documents analys\u00e9s : {n_docs}",
        f"Bilans : {n_bilans}  \u2014  Actes : {n_actes}",
        "Source : Registre National des Entreprises (INPI)",
    ]:
        story.append(RLPara(ml, sty["CvM"]))
        story.append(Spacer(1, 1.5 * mm))

    # ─── PAGE 2 : SOMMAIRE ────────────────────────────────────────────────
    story.append(NextPageTemplate('content'))
    story.append(PageBreak())
    story.append(RLPara("Sommaire", sty["TocT"]))
    story.append(HRFlowable(width="25%", thickness=3, color=C_ORANGE,
                             spaceBefore=0, spaceAfter=10 * mm, hAlign='LEFT'))

    toc = [
        (1, "A",   "Synth\u00e8se entreprise"),
        (2, "A.1", "Identit\u00e9"),
        (2, "A.2", "Objet social"),
        (2, "A.3", "Dirigeant"),
        (2, "A.4", "Adresse du si\u00e8ge"),
        (2, "A.5", "Activit\u00e9s (NAF)"),
        (2, "A.6", "Dates cl\u00e9s"),
        (2, "A.7", "Registres ant\u00e9rieurs"),
        (1, "B",   "Analyse des documents d\u00e9pos\u00e9s"),
        (2, "B.1", "Liste des documents"),
        (2, "B.2", "Analyse d\u00e9taill\u00e9e"),
        (1, "C",   "Points d\u2019attention"),
    ]
    for level, num, title in toc:
        s = sty["Toc1"] if level == 1 else sty["Toc2"]
        clr = _ohx if level == 1 else _mhx
        story.append(RLPara(
            f"<font color='#{clr}'><b>{num}</b></font>&nbsp;&nbsp;&nbsp;{_x(title)}", s))

    # ─── SECTION A : SYNTHÈSE ENTREPRISE ──────────────────────────────────
    story.append(PageBreak())
    story.append(RLPara(
        f"<font color='#{_ohx}'>A</font>&nbsp;&nbsp;Synth\u00e8se entreprise", sty["H1"]))
    story.append(HRFlowable(width="100%", thickness=1.5, color=C_ORANGE,
                             spaceBefore=0, spaceAfter=6 * mm))

    # A.1 Identité
    story.append(RLPara("A.1&nbsp;&nbsp;Identit\u00e9", sty["H2"]))
    story.append(_kv_table([
        ("D\u00e9nomination",  _dash(rne_data.get("denomination", denomination))),
        ("SIREN",              _dash(rne_data.get("siren", siren))),
        ("SIRET",              _dash(rne_data.get("siret"))),
        ("Code APE",           _dash(rne_data.get("ape"))),
        ("Forme juridique",    _dash(rne_data.get("forme"))),
        ("Capital social",     _dash(rne_data.get("capital"))),
        ("Dur\u00e9e",         _dash(rne_data.get("duree"))),
    ], sty, total_w=cw))

    # A.2 Objet social
    story.append(Spacer(1, 2 * mm))
    story.append(RLPara("A.2&nbsp;&nbsp;Objet social", sty["H2"]))
    objet = rne_data.get("objet")
    if objet:
        if isinstance(objet, list):
            story.append(_kv_table(
                [(f"Activit\u00e9 {i+1}", line) for i, line in enumerate(objet)],
                sty, total_w=cw))
        else:
            story.append(RLPara(_x(str(objet)), sty["Bd"]))
    else:
        story.append(RLPara("Non renseign\u00e9.", sty["Bd"]))

    # A.3 Dirigeant
    story.append(Spacer(1, 2 * mm))
    story.append(RLPara("A.3&nbsp;&nbsp;Dirigeant", sty["H2"]))
    story.append(_kv_table([
        ("Nom",               _dash(rne_data.get("dirigeant_nom"))),
        ("R\u00f4le",         _dash(rne_data.get("dirigeant_role"))),
        ("Date de naissance", _dash(rne_data.get("dirigeant_naissance"))),
    ], sty, total_w=cw))

    # A.4 Adresse du siège
    story.append(Spacer(1, 2 * mm))
    story.append(RLPara("A.4&nbsp;&nbsp;Adresse du si\u00e8ge", sty["H2"]))
    story.append(_kv_table([
        ("Adresse",              _dash(rne_data.get("adresse"))),
        ("Code INSEE commune",   _dash(rne_data.get("code_insee"))),
        ("Type de voie",         _dash(rne_data.get("type_voie"))),
    ], sty, total_w=cw))

    # A.5 Activités (NAF)
    story.append(Spacer(1, 2 * mm))
    story.append(RLPara("A.5&nbsp;&nbsp;Activit\u00e9s (NAF)", sty["H2"]))
    acts = rne_data.get("activites", [])
    if acts:
        if isinstance(acts[0], dict):
            hdr = [RLPara(f"<b>{h}</b>", sty["Th"]) for h in
                   ["Code APE", "Description", "Date de d\u00e9but"]]
            act_rows = [hdr]
            for a in acts:
                act_rows.append([
                    RLPara(_x(a.get("ape", "\u2014")), sty["Td"]),
                    RLPara(_x(a.get("desc", "\u2014")), sty["Td"]),
                    RLPara(_x(a.get("date", "\u2014")), sty["Td"]),
                ])
            at = Table(act_rows, colWidths=[4 * cm, cw - 7.5 * cm, 3.5 * cm])
            ats = [
                ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING",    (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING",   (0, 0), (-1, -1), 8),
                ("BACKGROUND",    (0, 0), (-1, 0), C_NAVY),
                ("TEXTCOLOR",     (0, 0), (-1, 0), white),
                ("LINEBELOW",     (0, 0), (-1, -1), 0.3, C_BORDER),
            ]
            for i in range(1, len(act_rows)):
                bg = C_LGRAY if i % 2 == 0 else C_CARD
                ats.append(("BACKGROUND", (0, i), (-1, i), bg))
            at.setStyle(TableStyle(ats))
            story.append(at)
        else:
            for a in acts:
                story.append(RLPara(f"\u2022&nbsp;&nbsp;{_x(str(a))}", sty["Bd"]))
    else:
        story.append(RLPara("Aucune activit\u00e9 enregistr\u00e9e.", sty["Bd"]))

    # A.6 Dates clés
    story.append(Spacer(1, 2 * mm))
    story.append(RLPara("A.6&nbsp;&nbsp;Dates cl\u00e9s", sty["H2"]))
    story.append(_kv_table([
        ("Date d\u2019immatriculation",          _dash(rne_data.get("date_immat"))),
        ("D\u00e9but d\u2019activit\u00e9",      _dash(rne_data.get("date_debut"))),
        ("Derni\u00e8re mise \u00e0 jour",       _dash(rne_data.get("date_maj"))),
        ("Derni\u00e8re modif. d\u2019activit\u00e9", _dash(rne_data.get("date_modif_activite"))),
        ("\u00c9tablissements ouverts",           _dash(rne_data.get("nb_etab"))),
        ("SIRET principal",                       _dash(rne_data.get("etab_siret", rne_data.get("siret")))),
    ], sty, total_w=cw))

    # A.7 Registres antérieurs
    story.append(Spacer(1, 2 * mm))
    story.append(RLPara("A.7&nbsp;&nbsp;Registres ant\u00e9rieurs", sty["H2"]))
    story.append(_kv_table([
        ("RNCS", _dash(rne_data.get("rncs"))),
        ("RNM",  _dash(rne_data.get("rnm"))),
        ("RAA",  _dash(rne_data.get("raa"))),
    ], sty, total_w=cw))

    # ─── SECTION B : DOCUMENTS DÉPOSÉS ────────────────────────────────────
    story.append(PageBreak())
    story.append(RLPara(
        f"<font color='#{_ohx}'>B</font>&nbsp;&nbsp;Analyse des documents d\u00e9pos\u00e9s",
        sty["H1"]))
    story.append(HRFlowable(width="100%", thickness=1.5, color=C_ORANGE,
                             spaceBefore=0, spaceAfter=6 * mm))

    # B.1 Liste des documents
    story.append(RLPara("B.1&nbsp;&nbsp;Liste des documents", sty["H2"]))
    if not doc_results:
        story.append(RLPara("Aucun document trouv\u00e9.", sty["Bd"]))
    else:
        hdr = [
            RLPara("<b>#</b>",             sty["Th"]),
            RLPara("<b>Date d\u00e9p\u00f4t</b>", sty["Th"]),
            RLPara("<b>Type</b>",           sty["Th"]),
            RLPara("<b>Nature INPI</b>",    sty["Th"]),
            RLPara("<b>Analys\u00e9</b>",   sty["Th"]),
        ]
        rows = [hdr]
        for ix, r in enumerate(doc_results, 1):
            analysed = r.get("document_analyse", "NON")
            a_sty = sty["BgO"] if analysed == "OUI" else sty["BgN"]
            td = r.get("typeDocument") or ""
            tl = r.get("typeLibelle") or r.get("nature") or "\u2014"
            nature = f"{td} \u2014 {tl}" if td else tl
            rows.append([
                RLPara(str(ix),                                  sty["Td"]),
                RLPara(_x(r.get("date_depot", "\u2014")),        sty["Td"]),
                RLPara(_x(r.get("famille", "\u2014")),           sty["Td"]),
                RLPara(_x(nature),                               sty["Td"]),
                RLPara(analysed,                                 a_sty),
            ])
        col_widths = [1 * cm, 3 * cm, 3.5 * cm, cw - 1 * cm - 3 * cm - 3.5 * cm - 2 * cm, 2 * cm]
        t = Table(rows, colWidths=col_widths)
        ts = [
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",    (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING",   (0, 0), (-1, -1), 6),
            ("BACKGROUND",    (0, 0), (-1, 0), C_NAVY),
            ("TEXTCOLOR",     (0, 0), (-1, 0), white),
            ("LINEBELOW",     (0, 1), (-1, -1), 0.3, C_BORDER),
        ]
        for i in range(1, len(rows)):
            bg = C_LGRAY if i % 2 == 0 else C_CARD
            ts.append(("BACKGROUND", (0, i), (-1, i), bg))
        t.setStyle(TableStyle(ts))
        story.append(t)

    # B.2 Analyse détaillée
    story.append(Spacer(1, 6 * mm))
    story.append(RLPara("B.2&nbsp;&nbsp;Analyse d\u00e9taill\u00e9e", sty["H2"]))

    for ix, r in enumerate(doc_results, 1):
        story.append(Spacer(1, 4 * mm))
        bdg = r.get("document_analyse", "NON")
        bhx = _ghx if bdg == "OUI" else _rhx
        fname = r.get("filename_base", r.get("filename", f"Document_{ix}"))
        story.append(RLPara(
            f"<b>Document {ix}/{n_docs}</b>&nbsp;&nbsp;"
            f"<font color='#{bhx}'>[{bdg}]</font>&nbsp;&nbsp;"
            f"{_x(fname)}.pdf",
            sty["DocT"]))
        td = r.get("typeDocument") or "N/A"
        tl = r.get("typeLibelle") or r.get("nature") or "N/A"
        detail_rows = [
            ("Date de d\u00e9p\u00f4t", r.get("date_depot", "\u2014")),
            ("Famille",                 r.get("famille", "\u2014")),
            ("Nature INPI",             f"{td} \u2014 {tl}"),
        ]
        story.append(_kv_table(detail_rows, sty, col1=4 * cm, total_w=cw))

        descriptif = r.get("descriptif") or r.get("description") or ""
        if descriptif:
            story.append(Spacer(1, 2 * mm))
            story.append(RLPara("<b>Descriptif</b>", sty["Lbl"]))
            for dl in _strip_md(descriptif.strip()).split("\n"):
                ds = dl.strip()
                if ds:
                    story.append(RLPara(_x(ds), sty["Bd"]))
            # Texte extrait
            texte = r.get("texte_extrait") or r.get("texte") or ""
            if texte:
                story.append(Spacer(1, 1 * mm))
                story.append(RLPara("<b>TEXTE EXTRAIT :</b>", sty["Lbl"]))
                excerpt = texte.strip()[:600]
                if len(texte.strip()) > 600:
                    excerpt += "…"
                story.append(RLPara(_x(excerpt), sty["BdS"]))
        else:
            story.append(RLPara("<i>Aucun descriptif disponible.</i>", sty["BdS"]))

        if ix < n_docs:
            story.append(Spacer(1, 3 * mm))
            story.append(HRFlowable(width="100%", thickness=0.3, color=C_BORDER,
                                     spaceBefore=0, spaceAfter=2 * mm))

    # ─── SECTION C : POINTS D'ATTENTION ───────────────────────────────────
    story.append(PageBreak())
    story.append(RLPara(
        f"<font color='#{_ohx}'>C</font>&nbsp;&nbsp;Points d\u2019attention", sty["H1"]))
    story.append(HRFlowable(width="100%", thickness=1.5, color=C_ORANGE,
                             spaceBefore=0, spaceAfter=6 * mm))

    pts = rne_data.get("points_attention", [])
    if pts:
        apars = [
            RLPara(f"<font color='#{_ohx}'>\u25b8</font>&nbsp;&nbsp;{_x(pt)}", sty["Alt"])
            for pt in pts
        ]
        story.append(AlertBox(cw, apars))
    else:
        story.append(RLPara("Aucun point d\u2019attention identifi\u00e9.", sty["Bd"]))

    # ─── PAGE FINALE (CLOSING) ─────────────────────────────────────────────
    story.append(NextPageTemplate('closing'))
    story.append(PageBreak())
    story.append(Spacer(1, 2.5 * cm))
    logo_cl = _logo_rl(5.5)
    if logo_cl:
        logo_cl.hAlign = 'CENTER'
        story.append(logo_cl)
    story.append(Spacer(1, 1.5 * cm))
    story.append(RLPara("Automatiser l\u2019analyse RNE dans votre cabinet", sty["ClT"]))
    story.append(Spacer(1, 6 * mm))
    story.append(RLPara(
        "Cet agent permet de collecter et structurer automatiquement les donn\u00e9es publiques "
        "du Registre National des Entreprises (INPI) afin de produire des rapports exploitables "
        "par les \u00e9quipes du cabinet.", sty["ClS"]))
    story.append(Spacer(1, 5 * mm))
    story.append(RLPara("Nous \u00e9tudions avec vous :", sty["ClS"]))
    story.append(Spacer(1, 2 * mm))
    for blt in [
        "l\u2019int\u00e9gration dans vos logiciels de production et de gestion interne",
        "l\u2019adaptation de l\u2019agent \u00e0 vos processus m\u00e9tier",
        "l\u2019automatisation de la collecte et de l\u2019analyse documentaire",
    ]:
        story.append(RLPara(f"\u2022  {blt}", sty["ClS"]))
    story.append(Spacer(1, 5 * mm))
    story.append(RLPara(
        "Les rapports g\u00e9n\u00e9r\u00e9s peuvent \u00e9galement \u00eatre personnalis\u00e9s "
        "aux couleurs de votre cabinet afin d\u2019\u00eatre utilis\u00e9s directement dans "
        "votre communication client.", sty["ClS"]))
    story.append(Spacer(1, 8 * mm))
    story.append(RLPara("<b>Contact : contact@iagidoo.fr</b>", sty["ClS"]))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# WORD — HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def _set_cell_bg(cell, hex_color):
    """Définit la couleur de fond d'une cellule de tableau Word."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.replace('#', ''))
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    """Définit les bordures d'une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        v = kwargs.get(side, {})
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), v.get('val', 'none'))
        border.set(qn('w:sz'), str(v.get('sz', 0)))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), v.get('color', 'auto'))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_heading1(doc, letter, title):
    """Ajoute un titre de section (ex : A  Synthèse entreprise)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run_letter = p.add_run(letter + "  ")
    run_letter.font.color.rgb = D_ORANGE
    run_letter.font.size = Pt(20)
    run_letter.font.bold = True
    run_title = p.add_run(title)
    run_title.font.color.rgb = D_NAVY
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    # Ligne de séparation orange
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(8)
    run_hr = p2.add_run("_" * 90)
    run_hr.font.color.rgb = D_ORANGE
    run_hr.font.size = Pt(5)


def _add_heading2(doc, num, title):
    """Ajoute un sous-titre de section (ex : A.1  Identité)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run_num = p.add_run(num + "  ")
    run_num.font.color.rgb = D_TEAL
    run_num.font.size = Pt(13)
    run_num.font.bold = True
    run_title = p.add_run(title)
    run_title.font.color.rgb = D_TEAL
    run_title.font.size = Pt(13)
    run_title.font.bold = True


def _add_kv_table(doc, rows, col1_cm=5.5, total_cm=17.0):
    """Ajoute une table clé/valeur avec lignes alternées."""
    col2_cm = total_cm - col1_cm
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Cm(col1_cm)
    table.columns[1].width = Cm(col2_cm)

    for i, (lbl, val) in enumerate(rows):
        row = table.rows[i]
        bg = "F3F1EE" if i % 2 == 0 else "FFFFFF"

        # Colonne label
        cell0 = row.cells[0]
        _set_cell_bg(cell0, bg)
        p0 = cell0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after = Pt(3)
        run0 = p0.add_run(str(lbl))
        run0.font.bold = True
        run0.font.size = Pt(9)
        run0.font.color.rgb = D_MUTED

        # Colonne valeur
        cell1 = row.cells[1]
        _set_cell_bg(cell1, bg)
        p1 = cell1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after = Pt(3)
        run1 = p1.add_run(str(val))
        run1.font.size = Pt(9.5)

    # Supprimer les bordures extérieures, garder seulement les lignes intérieures légères
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(
                cell,
                bottom={'val': 'single', 'sz': 2, 'color': 'E8E4DF'},
                top={'val': 'none'},
                left={'val': 'none'},
                right={'val': 'none'},
            )
    return table


def _add_doc_table_header(doc, headers, col_widths_cm):
    """Ajoute une table avec en-tête navy."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (h, w) in enumerate(zip(headers, col_widths_cm)):
        table.columns[i].width = Cm(w)
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, "1A2744")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = D_WHITE
    return table


def _add_doc_table_row(table, values, row_index):
    """Ajoute une ligne à un tableau de documents."""
    row = table.add_row()
    bg = "F3F1EE" if row_index % 2 == 0 else "FFFFFF"
    for i, val in enumerate(values):
        cell = row.cells[i]
        _set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if isinstance(val, tuple):
            text, color = val
            run = p.add_run(str(text))
            run.font.size = Pt(8.5)
            run.font.bold = True
            if color == "green":
                run.font.color.rgb = D_GREEN
            elif color == "red":
                run.font.color.rgb = D_RED
        else:
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
        _set_cell_border(
            cell,
            bottom={'val': 'single', 'sz': 2, 'color': 'E8E4DF'},
            top={'val': 'none'},
            left={'val': 'none'},
            right={'val': 'none'},
        )


def _add_alert_box(doc, points):
    """Ajoute le bloc 'Points d'attention' avec fond orange."""
    if not points:
        p = doc.add_paragraph("Aucun point d'attention identifié.")
        p.runs[0].font.size = Pt(9.5)
        return

    # Titre du bloc
    table = doc.add_table(rows=1 + len(points), cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Cm(17.0)

    # En-tête orange
    cell_hdr = table.rows[0].cells[0]
    _set_cell_bg(cell_hdr, "EF8829")
    p_hdr = cell_hdr.paragraphs[0]
    p_hdr.paragraph_format.space_before = Pt(4)
    p_hdr.paragraph_format.space_after = Pt(4)
    run_hdr = p_hdr.add_run("Points d'attention")
    run_hdr.font.bold = True
    run_hdr.font.size = Pt(9.5)
    run_hdr.font.color.rgb = D_WHITE

    # Lignes des points
    for i, pt in enumerate(points):
        cell = table.rows[i + 1].cells[0]
        _set_cell_bg(cell, "FFF8F0")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.3)
        run_bullet = p.add_run("\u25b8  ")
        run_bullet.font.color.rgb = D_ORANGE
        run_bullet.font.size = Pt(9.5)
        run_text = p.add_run(str(pt))
        run_text.font.size = Pt(9.5)
        _set_cell_border(
            cell,
            top={'val': 'none'},
            left={'val': 'single', 'sz': 4, 'color': 'EF8829'},
            right={'val': 'single', 'sz': 4, 'color': 'EF8829'},
            bottom={'val': 'single', 'sz': 2, 'color': 'E8E4DF'} if i < len(points) - 1
                  else {'val': 'single', 'sz': 4, 'color': 'EF8829'},
        )


# ══════════════════════════════════════════════════════════════════════════════
# GÉNÉRATION WORD (DOCX)
# ══════════════════════════════════════════════════════════════════════════════
def generate_report_docx(siren, denomination, rne_data, doc_results, run_date):
    """
    Génère le rapport d'analyse RNE au format Word (DOCX).

    Returns:
        bytes: Contenu du fichier DOCX
    """
    doc = Document()

    # ── Marges de page ──
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    n_docs   = len(doc_results)
    n_bilans = sum(1 for d in doc_results if (d.get("famille") or "").upper().startswith("COMP"))
    n_actes  = sum(1 for d in doc_results if (d.get("famille") or "").upper() == "ACTE")

    # ─── PAGE 1 : COUVERTURE ──────────────────────────────────────────────
    # Logo (si disponible)
    if os.path.exists(LOGO_PATH):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(LOGO_PATH, width=Cm(6))
    else:
        p_logo = doc.add_paragraph("GIDOO")
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.runs[0].font.size = Pt(28)
        p_logo.runs[0].font.bold = True

    doc.add_paragraph()  # espace

    # Titre
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_t = p_title.add_run("Rapport d\u2019analyse RNE")
    run_t.font.size = Pt(32)
    run_t.font.bold = True
    run_t.font.color.rgb = D_NAVY

    # Sous-titre
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(denomination)
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = D_TEAL

    # SIREN
    p_siren = doc.add_paragraph()
    run_sr = p_siren.add_run(f"SIREN {siren}")
    run_sr.font.size = Pt(11)
    run_sr.font.color.rgb = D_MUTED

    # Séparateur
    p_sep = doc.add_paragraph("_" * 90)
    p_sep.runs[0].font.color.rgb = D_ORANGE
    p_sep.runs[0].font.size = Pt(5)

    # Métadonnées
    for ml in [
        f"Date du rapport : {run_date}",
        f"Documents analys\u00e9s : {n_docs}",
        f"Bilans : {n_bilans}   \u2014   Actes : {n_actes}",
        "Source : Registre National des Entreprises (INPI)",
    ]:
        p_m = doc.add_paragraph()
        run_m = p_m.add_run(ml)
        run_m.font.size = Pt(10)
        run_m.font.color.rgb = D_MUTED

    # ─── PAGE 2 : SOMMAIRE ────────────────────────────────────────────────
    doc.add_page_break()

    p_toc_title = doc.add_paragraph()
    run_toc = p_toc_title.add_run("Sommaire")
    run_toc.font.size = Pt(22)
    run_toc.font.bold = True
    run_toc.font.color.rgb = D_NAVY
    p_toc_title.paragraph_format.space_after = Pt(8)

    p_sep2 = doc.add_paragraph("_" * 30)
    p_sep2.runs[0].font.color.rgb = D_ORANGE
    p_sep2.runs[0].font.size = Pt(5)
    p_sep2.paragraph_format.space_after = Pt(12)

    toc = [
        (1, "A",   "Synth\u00e8se entreprise"),
        (2, "A.1", "Identit\u00e9"),
        (2, "A.2", "Objet social"),
        (2, "A.3", "Dirigeant"),
        (2, "A.4", "Adresse du si\u00e8ge"),
        (2, "A.5", "Activit\u00e9s (NAF)"),
        (2, "A.6", "Dates cl\u00e9s"),
        (2, "A.7", "Registres ant\u00e9rieurs"),
        (1, "B",   "Analyse des documents d\u00e9pos\u00e9s"),
        (2, "B.1", "Liste des documents"),
        (2, "B.2", "Analyse d\u00e9taill\u00e9e"),
        (1, "C",   "Points d\u2019attention"),
    ]
    for level, num, title in toc:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_before = Pt(0)
        p_item.paragraph_format.space_after = Pt(0)
        if level == 2:
            p_item.paragraph_format.left_indent = Cm(0.8)
        run_num = p_item.add_run(f"{num}   ")
        run_num.font.bold = True
        run_num.font.size = Pt(11) if level == 1 else Pt(10)
        run_num.font.color.rgb = D_ORANGE if level == 1 else D_MUTED
        run_ttl = p_item.add_run(title)
        run_ttl.font.size = Pt(11) if level == 1 else Pt(10)
        run_ttl.font.bold = (level == 1)

    # ─── SECTION A : SYNTHÈSE ENTREPRISE ──────────────────────────────────
    doc.add_page_break()
    _add_heading1(doc, "A", "Synth\u00e8se entreprise")

    _add_heading2(doc, "A.1", "Identit\u00e9")
    _add_kv_table(doc, [
        ("D\u00e9nomination",  _dash(rne_data.get("denomination", denomination))),
        ("SIREN",              _dash(rne_data.get("siren", siren))),
        ("SIRET",              _dash(rne_data.get("siret"))),
        ("Code APE",           _dash(rne_data.get("ape"))),
        ("Forme juridique",    _dash(rne_data.get("forme"))),
        ("Capital social",     _dash(rne_data.get("capital"))),
        ("Dur\u00e9e",         _dash(rne_data.get("duree"))),
    ])

    _add_heading2(doc, "A.2", "Objet social")
    objet = rne_data.get("objet")
    if objet:
        if isinstance(objet, list):
            _add_kv_table(doc, [(f"Activit\u00e9 {i+1}", line) for i, line in enumerate(objet)])
        else:
            p = doc.add_paragraph(str(objet))
            p.runs[0].font.size = Pt(9.5)
    else:
        p = doc.add_paragraph("Non renseign\u00e9.")
        p.runs[0].font.size = Pt(9.5)

    _add_heading2(doc, "A.3", "Dirigeant")
    _add_kv_table(doc, [
        ("Nom",               _dash(rne_data.get("dirigeant_nom"))),
        ("R\u00f4le",         _dash(rne_data.get("dirigeant_role"))),
        ("Date de naissance", _dash(rne_data.get("dirigeant_naissance"))),
    ])

    _add_heading2(doc, "A.4", "Adresse du si\u00e8ge")
    _add_kv_table(doc, [
        ("Adresse",            _dash(rne_data.get("adresse"))),
        ("Code INSEE commune", _dash(rne_data.get("code_insee"))),
        ("Type de voie",       _dash(rne_data.get("type_voie"))),
    ])

    _add_heading2(doc, "A.5", "Activit\u00e9s (NAF)")
    acts = rne_data.get("activites", [])
    if acts:
        if isinstance(acts[0], dict):
            headers = ["Code APE", "Description", "Date de d\u00e9but"]
            widths  = [4.0, 9.0, 4.0]
            table_acts = _add_doc_table_header(doc, headers, widths)
            for i, a in enumerate(acts):
                _add_doc_table_row(
                    table_acts,
                    [a.get("ape", "\u2014"), a.get("desc", "\u2014"), a.get("date", "\u2014")],
                    i,
                )
        else:
            for a in acts:
                p = doc.add_paragraph(f"\u2022  {a}")
                p.runs[0].font.size = Pt(9.5)
    else:
        p = doc.add_paragraph("Aucune activit\u00e9 enregistr\u00e9e.")
        p.runs[0].font.size = Pt(9.5)

    _add_heading2(doc, "A.6", "Dates cl\u00e9s")
    _add_kv_table(doc, [
        ("Date d\u2019immatriculation",          _dash(rne_data.get("date_immat"))),
        ("D\u00e9but d\u2019activit\u00e9",      _dash(rne_data.get("date_debut"))),
        ("Derni\u00e8re mise \u00e0 jour",       _dash(rne_data.get("date_maj"))),
        ("Derni\u00e8re modif. d\u2019activit\u00e9", _dash(rne_data.get("date_modif_activite"))),
        ("\u00c9tablissements ouverts",           _dash(rne_data.get("nb_etab"))),
        ("SIRET principal",                       _dash(rne_data.get("etab_siret", rne_data.get("siret")))),
    ])

    _add_heading2(doc, "A.7", "Registres ant\u00e9rieurs")
    _add_kv_table(doc, [
        ("RNCS", _dash(rne_data.get("rncs"))),
        ("RNM",  _dash(rne_data.get("rnm"))),
        ("RAA",  _dash(rne_data.get("raa"))),
    ])

    # ─── SECTION B : DOCUMENTS DÉPOSÉS ────────────────────────────────────
    doc.add_page_break()
    _add_heading1(doc, "B", "Analyse des documents d\u00e9pos\u00e9s")

    _add_heading2(doc, "B.1", "Liste des documents")
    if not doc_results:
        p = doc.add_paragraph("Aucun document trouv\u00e9.")
        p.runs[0].font.size = Pt(9.5)
    else:
        headers = ["#", "Date d\u00e9p\u00f4t", "Type", "Nature INPI", "Analys\u00e9"]
        widths  = [1.0, 2.8, 3.2, 8.0, 2.0]
        table_docs = _add_doc_table_header(doc, headers, widths)
        for ix, r in enumerate(doc_results, 1):
            analysed = r.get("document_analyse", "NON")
            td = r.get("typeDocument") or ""
            tl = r.get("typeLibelle") or r.get("nature") or "\u2014"
            nature = f"{td} \u2014 {tl}" if td else tl
            badge_color = "green" if analysed == "OUI" else "red"
            _add_doc_table_row(
                table_docs,
                [str(ix), r.get("date_depot", "\u2014"), r.get("famille", "\u2014"),
                 nature, (analysed, badge_color)],
                ix - 1,
            )

    _add_heading2(doc, "B.2", "Analyse d\u00e9taill\u00e9e")
    for ix, r in enumerate(doc_results, 1):
        doc.add_paragraph()
        bdg = r.get("document_analyse", "NON")
        fname = r.get("filename_base", r.get("filename", f"Document_{ix}"))

        p_doc_title = doc.add_paragraph()
        p_doc_title.paragraph_format.space_before = Pt(6)
        run_num  = p_doc_title.add_run(f"Document {ix}/{n_docs}  ")
        run_num.font.bold = True
        run_num.font.size = Pt(11)
        run_num.font.color.rgb = D_NAVY
        run_badge = p_doc_title.add_run(f"[{bdg}]  ")
        run_badge.font.bold = True
        run_badge.font.size = Pt(11)
        run_badge.font.color.rgb = D_GREEN if bdg == "OUI" else D_RED
        run_name = p_doc_title.add_run(f"{fname}.pdf")
        run_name.font.size = Pt(11)
        run_name.font.color.rgb = D_NAVY

        td = r.get("typeDocument") or "N/A"
        tl = r.get("typeLibelle") or r.get("nature") or "N/A"
        _add_kv_table(doc, [
            ("Date de d\u00e9p\u00f4t", r.get("date_depot", "\u2014")),
            ("Famille",                 r.get("famille", "\u2014")),
            ("Nature INPI",             f"{td} \u2014 {tl}"),
        ], col1_cm=4.0)

        descriptif = r.get("descriptif") or r.get("description") or ""
        if descriptif:
            p_lbl = doc.add_paragraph()
            p_lbl.paragraph_format.space_before = Pt(4)
            run_lbl = p_lbl.add_run("Descriptif")
            run_lbl.font.bold = True
            run_lbl.font.size = Pt(9)
            run_lbl.font.color.rgb = D_MUTED

            for dl in _strip_md(descriptif.strip()).split("\n"):
                ds = dl.strip()
                if ds:
                    p_d = doc.add_paragraph(ds)
                    p_d.runs[0].font.size = Pt(9.5)

            texte = r.get("texte_extrait") or r.get("texte") or ""
            if texte:
                p_te_lbl = doc.add_paragraph()
                run_te = p_te_lbl.add_run("TEXTE EXTRAIT :")
                run_te.font.bold = True
                run_te.font.size = Pt(9)
                run_te.font.color.rgb = D_MUTED
                excerpt = texte.strip()[:600]
                if len(texte.strip()) > 600:
                    excerpt += "…"
                p_te = doc.add_paragraph(excerpt)
                p_te.runs[0].font.size = Pt(8.5)
                p_te.runs[0].font.color.rgb = D_MUTED
        else:
            p_nd = doc.add_paragraph("Aucun descriptif disponible.")
            p_nd.runs[0].font.italic = True
            p_nd.runs[0].font.size = Pt(8.5)
            p_nd.runs[0].font.color.rgb = D_MUTED

        if ix < n_docs:
            p_sep_d = doc.add_paragraph("_" * 90)
            p_sep_d.runs[0].font.color.rgb = RGBColor(0xE8, 0xE4, 0xDF)
            p_sep_d.runs[0].font.size = Pt(4)

    # ─── SECTION C : POINTS D'ATTENTION ───────────────────────────────────
    doc.add_page_break()
    _add_heading1(doc, "C", "Points d\u2019attention")
    _add_alert_box(doc, rne_data.get("points_attention", []))

    # ─── PAGE FINALE ──────────────────────────────────────────────────────
    doc.add_page_break()

    # Fond simulé via paragraphes de couleur
    if os.path.exists(LOGO_PATH):
        p_logo2 = doc.add_paragraph()
        p_logo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo2.add_run().add_picture(LOGO_PATH, width=Cm(6))

    p_cl = doc.add_paragraph()
    p_cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cl.paragraph_format.space_before = Pt(20)
    run_cl = p_cl.add_run("Automatiser l\u2019analyse RNE dans votre cabinet")
    run_cl.font.size = Pt(16)
    run_cl.font.bold = True
    run_cl.font.color.rgb = D_NAVY

    p_cl2 = doc.add_paragraph()
    p_cl2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cl2 = p_cl2.add_run(
        "Cet agent permet de collecter et structurer automatiquement les donn\u00e9es publiques "
        "du Registre National des Entreprises (INPI) afin de produire des rapports exploitables "
        "par les \u00e9quipes du cabinet.")
    run_cl2.font.size = Pt(10)
    run_cl2.font.color.rgb = D_MUTED

    p_nous = doc.add_paragraph()
    p_nous.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nous.paragraph_format.space_before = Pt(10)
    run_nous = p_nous.add_run("Nous \u00e9tudions avec vous :")
    run_nous.font.size = Pt(10)
    run_nous.font.color.rgb = D_MUTED

    for blt in [
        "l\u2019int\u00e9gration dans vos logiciels de production et de gestion interne",
        "l\u2019adaptation de l\u2019agent \u00e0 vos processus m\u00e9tier",
        "l\u2019automatisation de la collecte et de l\u2019analyse documentaire",
    ]:
        p_b = doc.add_paragraph(f"\u2022  {blt}")
        p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_b.runs[0].font.size = Pt(10)
        p_b.runs[0].font.color.rgb = D_MUTED

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(14)
    run_contact = p_contact.add_run("Contact : contact@iagidoo.fr")
    run_contact.font.size = Pt(11)
    run_contact.font.bold = True
    run_contact.font.color.rgb = D_TEAL

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# POINT D'ENTRÉE
# ══════════════════════════════════════════════════════════════════════════════
def _sanitize(name):
    return re.sub(r'\s+', '_', re.sub(r'[\\/:*?"<>|;]', '', name)).strip('_')


def generate_reports(siren, denomination, rne_data, doc_results,
                     output_dir=None, run_date=None):
    """
    Génère les fichiers PDF et DOCX du rapport RNE.

    Returns:
        tuple[str, str]: chemins des fichiers PDF et DOCX générés
    """
    if run_date is None:
        run_date = datetime.today().strftime("%Y-%m-%d")
    if output_dir is None:
        output_dir = os.path.dirname(os.path.abspath(__file__))

    base = _sanitize(f"{run_date}_{siren}_{denomination}_Rapport_Analyse_RNE")
    pdf_path  = os.path.join(output_dir, base + ".pdf")
    docx_path = os.path.join(output_dir, base + ".docx")

    pdf_bytes  = generate_report_pdf(siren, denomination, rne_data, doc_results, run_date)
    docx_bytes = generate_report_docx(siren, denomination, rne_data, doc_results, run_date)

    with open(pdf_path, "wb") as f:
        f.write(pdf_bytes)
    with open(docx_path, "wb") as f:
        f.write(docx_bytes)

    print(f"[OK] PDF  : {pdf_path}")
    print(f"[OK] DOCX : {docx_path}")
    return pdf_path, docx_path


# ══════════════════════════════════════════════════════════════════════════════
# EXEMPLE AVEC LES DONNÉES DU RAPPORT JOINT
# ══════════════════════════════════════════════════════════════════════════════
EXAMPLE_RNE_DATA = {
    "denomination":         "PEDA GO EXPERTS",
    "siren":                "983670290",
    "siret":                "98367029000011",
    "ape":                  "7022Z (Conseil pour les affaires et autres conseils de gestion)",
    "forme":                "SASU (5710) (Société par actions simplifiée unipersonnelle)",
    "capital":              "1 000 € (fixe, non variable)",
    "duree":                "99 ans (clôture exercice social au 31/12)",
    "objet":                None,
    "dirigeant_nom":        "PEDA GO EXPERTS",
    "dirigeant_role":       "Gérant(e) unique (associé unique et dirigeant)",
    "dirigeant_naissance":  "1987-03 (La Possession, 97419)",
    "adresse":              "ZA Ravine à Marquet, Rue Antanifotsy, 97419 La Possession",
    "code_insee":           None,
    "type_voie":            None,
    "activites":            ["7022Z (Conseil pour les affaires et autres conseils de gestion)"],
    "date_immat":           "19/01/2024",
    "date_debut":           None,
    "date_maj":             "03/07/2024",
    "date_modif_activite":  None,
    "nb_etab":              None,
    "etab_siret":           "98367029000011",
    "rncs":                 None,
    "rnm":                  None,
    "raa":                  None,
    "points_attention": [
        "Aucun bilan déposé (entreprise récente, pas d'obligation comptable encore).",
        "Siège non validé par l'INSEE (risque de non-conformité).",
        "Aucun établissement secondaire déclaré, alors que l'indicateur le suggère.",
        "Capital social faible (1 000 €), cohérent avec une SASU.",
    ],
}

EXAMPLE_DOC_RESULTS = [
    {
        "filename_base":    "2024-01-19 - 983670290 - Copie des statuts - ACTE",
        "date_depot":       "2024-01-19",
        "famille":          "ACTE",
        "typeDocument":     "PJ_01",
        "typeLibelle":      "Copie des statuts",
        "document_analyse": "NON",
        "descriptif": (
            "Le document contient les statuts d'une société, incluant les clauses relatives "
            "à la dénomination sociale, l'objet social, le siège social, les modalités de "
            "fonctionnement, les règles de gouvernance (assemblée générale, dirigeants), "
            "ainsi que les dispositions relatives à la dissolution et à la liquidation. "
            "Les statuts précisent également les droits et obligations des associés ou "
            "actionnaires, ainsi que les conditions de modification des statuts."
        ),
        "texte_extrait": "[Texte des statuts complet, incluant les articles et clauses mentionnés ci-dessus.]",
    },
    {
        "filename_base":    "2024-01-19 - 983670290 - - - Certificat dépositaire fonds",
        "date_depot":       "2024-01-19",
        "famille":          "ACTE",
        "typeDocument":     "PJ_06",
        "typeLibelle":      "Certificat du dépositaire des fonds",
        "document_analyse": "OUI",
        "descriptif": (
            "Le document certifie que les fonds ont été déposés auprès d'un dépositaire "
            "agréé, conformément aux dispositions légales en vigueur. Il mentionne la date "
            "de dépôt, le montant versé et les références du compte de dépôt. Le certificat "
            "est signé par le dépositaire et porte son cachet officiel."
        ),
    },
]


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Génère un rapport RNE PDF + DOCX")
    parser.add_argument("--siren",  default="983670290", help="Numéro SIREN")
    parser.add_argument("--output", default=None,        help="Dossier de sortie")
    args = parser.parse_args()

    # Utilise les données d'exemple si le SIREN correspond
    if args.siren == "983670290":
        rne_data    = EXAMPLE_RNE_DATA
        doc_results = EXAMPLE_DOC_RESULTS
        denomination = rne_data["denomination"]
    else:
        print(f"Pas de données d'exemple pour le SIREN {args.siren}.")
        print("Renseignez les variables rne_data et doc_results manuellement.")
        exit(1)

    generate_reports(
        siren=args.siren,
        denomination=denomination,
        rne_data=rne_data,
        doc_results=doc_results,
        output_dir=args.output,
    )
